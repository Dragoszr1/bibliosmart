from flask import Blueprint, jsonify, request, send_from_directory, current_app, g, send_file
from sqlalchemy import text
from sqlalchemy.exc import IntegrityError
from functools import wraps
import bcrypt
import jwt
import datetime
import io
import os
import logging
import re
import secrets
import random
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from groq import Groq
import smtplib
from email.mime.text import MIMEText

from app.database import db
from app.models import Carti, Users, CartiCitite, Recenzii, Anunturi, AnunturiAprecieri, CereriCarti, ClubThreads, ThreadComments, ThreadSubcomments

main_bp = Blueprint('main', __name__, url_prefix='/api')

ALLOWED_EXTENSIONS = frozenset({'png', 'jpg', 'jpeg', 'gif', 'webp'})
ALLOWED_EMAIL_DOMAIN = 'cni-sv.ro'
VALID_ROLES = frozenset({'user', 'bibliotecar'})
ALLOWED_BOOK_FIELDS = frozenset({'titlu', 'autor', 'ISBN', 'stoc_total', 'stoc_disponibil', 'gen', 'pozitie', 'cod'})
ALLOWED_ANUNT_FIELDS = frozenset({'titlu', 'anunt'})

logger = logging.getLogger(__name__)
ROLE_ALIASES = {
    '1': 'bibliotecar',
    'administrator': 'bibliotecar',
    'admin': 'bibliotecar',
    'bibliotecar': 'bibliotecar',
    'user': 'user'
}
EMAIL_REGEX = re.compile(r"^[A-Za-z0-9](?:[A-Za-z0-9._%+\-]{0,62}[A-Za-z0-9])?@([A-Za-z0-9\-]+\.)+[A-Za-z]{2,63}$")
USERNAME_REGEX = re.compile(r"^[A-Za-z0-9ăîâșțĂÎÂȘȚ\s._-]{3,50}$")
PROFILE_PICTURES_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), 'profile_pictures')
BOOK_IMAGES_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), 'book_images')
BOOK_PDFS_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), 'book_pdfs')
CLUB_ANNOUNCEMENT_IMAGES_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), 'imagini_anunturi_club')

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def _get_request_data():
    if request.content_type and request.content_type.startswith('multipart/form-data'):
        return request.form.to_dict()
    return request.get_json(silent=True) or {}


def _save_club_activity_image(file, activitate_id):
    if not file or not file.filename:
        return None
    if not allowed_file(file.filename):
        return None

    os.makedirs(CLUB_ANNOUNCEMENT_IMAGES_DIR, exist_ok=True)
    for old in _find_files_by_prefix(CLUB_ANNOUNCEMENT_IMAGES_DIR, str(activitate_id)):
        try:
            os.remove(os.path.join(CLUB_ANNOUNCEMENT_IMAGES_DIR, old))
        except OSError:
            pass

    ext = file.filename.rsplit('.', 1)[1].lower()
    filename = f"{activitate_id}.{ext}"
    filepath = os.path.join(CLUB_ANNOUNCEMENT_IMAGES_DIR, filename)
    file.save(filepath)
    return filename


def _find_files_by_prefix(directory, prefix):
    """Caută fișiere în director cu prefixul dat și extensie permisă.
    Folosește os.listdir în loc de glob ca să evite glob injection din
    șiruri controlate de utilizator care conțin *, ?, sau [].
    """
    os.makedirs(directory, exist_ok=True)
    rezultate = []
    for fisier in os.listdir(directory):
        if not fisier.startswith(prefix + '.'):
            continue
        ext = fisier.rsplit('.', 1)[1].lower() if '.' in fisier else ''
        if ext in ALLOWED_EXTENSIONS:
            rezultate.append(fisier)
    return rezultate


# email

def send_email(recipients, subject, body):
    """Trimite un email plain-text prin SMTP (Gmail).

    Citește setările SMTP_* din config-ul Flask.
    Returnează True la succes, False la eroare (logat, nu aruncă excepție).
    """
    cfg = current_app.config
    host = cfg.get('SMTP_HOST', 'smtp.gmail.com')
    port = cfg.get('SMTP_PORT', 587)
    user = cfg.get('SMTP_USER', '')
    password = cfg.get('SMTP_PASSWORD', '')
    sender = cfg.get('SMTP_FROM', '') or user

    if not user or not password:
        logger.error('Credențiale SMTP lipsă — emailul nu a fost trimis')
        return False

    msg = MIMEText(body, 'plain', 'utf-8')
    msg['Subject'] = subject
    msg['From'] = sender
    msg['To'] = ', '.join(recipients)

    try:
        with smtplib.SMTP(host, port, timeout=10) as server:
            server.ehlo()
            server.starttls()
            server.ehlo()
            server.login(user, password)
            server.sendmail(sender, recipients, msg.as_string())
        logger.info('Email trimis către %s', recipients)
        return True
    except Exception:
        logger.exception('Eroare la trimiterea emailului')
        return False


# jwt

def normalize_role(raw_role):
    if raw_role is None:
        return 'user'

    rol_normalizat = ROLE_ALIASES.get(str(raw_role).strip().lower())
    return rol_normalizat if rol_normalizat in VALID_ROLES else 'user'


def normalize_cni_email(email):
    if not isinstance(email, str):
        return None

    email_normalizat = email.strip().lower()
    if not email_normalizat or len(email_normalizat) > 254:
        return None

    if not EMAIL_REGEX.fullmatch(email_normalizat):
        return None

    parte_locala, _, domeniu = email_normalizat.rpartition('@')
    if not parte_locala or '..' in parte_locala or domeniu != ALLOWED_EMAIL_DOMAIN:
        return None

    return email_normalizat


def fetch_user_by_id(user_id):
    query = text(
        "SELECT user_id, username, email, rol, club FROM users WHERE user_id = :user_id"
    )
    result = db.session.execute(query, {'user_id': user_id}).mappings().first()
    if not result:
        return None

    return {
        'user_id': result['user_id'],
        'username': result['username'],
        'email': result['email'],
        'rol': normalize_role(result['rol']),
        'club': bool(result['club'])
    }


def create_jwt_token(user_id, username, email):
    payload = {
        'user_id': user_id,
        'username': username,
        'email': email,
        'iat': datetime.datetime.now(datetime.timezone.utc),
        'exp': datetime.datetime.now(datetime.timezone.utc) + current_app.config['JWT_ACCESS_TOKEN_EXPIRES']
    }
    return jwt.encode(payload, current_app.config['JWT_SECRET_KEY'], algorithm='HS256')


def set_jwt_cookie(response, token):
    response.set_cookie(
        current_app.config['JWT_COOKIE_NAME'],
        token,
        httponly=True,
        secure=current_app.config['JWT_COOKIE_SECURE'],
        samesite=current_app.config['JWT_COOKIE_SAMESITE'],
        max_age=int(current_app.config['JWT_ACCESS_TOKEN_EXPIRES'].total_seconds()),
        path='/'
    )
    return response


def clear_jwt_cookie(response):
    response.set_cookie(
        current_app.config['JWT_COOKIE_NAME'],
        '',
        httponly=True,
        secure=current_app.config['JWT_COOKIE_SECURE'],
        samesite=current_app.config['JWT_COOKIE_SAMESITE'],
        max_age=0,
        path='/'
    )
    return response


def decode_jwt_token(token):
    try:
        return jwt.decode(token, current_app.config['JWT_SECRET_KEY'], algorithms=['HS256'])
    except (jwt.ExpiredSignatureError, jwt.InvalidTokenError):
        return None


def get_current_user():
    """Rezolvă utilizatorul autentificat din cookie-ul JWT și baza de date.
    Cachează rezultatul pe flask.g ca să nu se mai facă query la DB
    pentru apeluri repetate în aceeași cerere.
    """
    if hasattr(g, '_current_user'):
        return g._current_user

    token = request.cookies.get(current_app.config['JWT_COOKIE_NAME'])
    if not token:
        g._current_user = None
        return None

    payload = decode_jwt_token(token)
    if not payload:
        g._current_user = None
        return None

    try:
        user_id = int(payload.get('user_id'))
    except (TypeError, ValueError):
        g._current_user = None
        return None

    g._current_user = fetch_user_by_id(user_id)
    return g._current_user


def jwt_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        user = get_current_user()
        if not user:
            return jsonify({'success': False, 'message': 'Autentificare necesară'}), 401
        request.current_user = user
        return f(*args, **kwargs)
    return decorated


def role_required(role):
    if role not in VALID_ROLES:
        raise ValueError(f"Rol necunoscut: {role}")

    def decorator(f):
        @wraps(f)
        @jwt_required
        def decorated(*args, **kwargs):
            if request.current_user.get('rol') != role:
                return jsonify({'success': False, 'message': f'Necesită rol de {role}'}), 403
            return f(*args, **kwargs)
        return decorated
    return decorator


bibliotecar_required = role_required('bibliotecar')

@main_bp.route('/health', methods=['GET'])
def health_check():
    return jsonify({'status': 'ok', 'message': 'Server is running'}), 200

@main_bp.route('/', methods=['GET'])
def index():
    return jsonify({
        'name': 'API site biblioteca CNI Suceava',
        'version': 'idk(forgot to count lol)',
        'message': 'Database connection established'
    }), 200

@main_bp.route('/books', methods=['GET'])
def get_books():
    try:
        result = db.session.execute(text('SELECT carte_id, titlu, autor, ISBN, stoc_total, stoc_disponibil, imprumutat, gen, pozitie, cod FROM carti'))
        os.makedirs(BOOK_PDFS_DIR, exist_ok=True)
        pdf_ids = {f.rsplit('.', 1)[0] for f in os.listdir(BOOK_PDFS_DIR) if f.endswith('.pdf')}
        books = []
        for row in result:
            books.append({
                'carte_id': row[0],
                'titlu': row[1],
                'autor': row[2],
                'ISBN': row[3],
                'stoc_total': row[4],
                'stoc_disponibil': row[5],
                'imprumutat': row[6],
                'gen': row[7],
                'pozitie': row[8],
                'cod': row[9],
                'has_pdf': str(row[0]) in pdf_ids
            })
        return jsonify({'books': books}), 200
    except Exception:
        logger.exception('Eroare la preluarea cărților')
        return jsonify({'error': 'Eroare la preluarea cărților'}), 500


@main_bp.route('/books/recent', methods=['GET'])
def get_recent_books():
    try:
        # Fetch the 3 books with the highest ID (assumes carte_id auto-increments and correlates with creation time)
        result = db.session.execute(text('SELECT carte_id, titlu, autor, gen, stoc_disponibil FROM carti ORDER BY carte_id DESC LIMIT 3'))
        
        books = []
        for row in result:
            books.append({
                'carte_id': row[0],
                'titlu': row[1],
                'autor': row[2],
                'gen': row[3],
                'stoc_disponibil': row[4]
            })
            
        return jsonify({'success': True, 'books': books}), 200
    except Exception:
        logger.exception('Eroare la preluarea cărților recente')
        return jsonify({'success': False, 'message': 'Eroare la preluarea cărților recente'}), 500


@main_bp.route('/books', methods=['POST'])
@bibliotecar_required
def add_book():
    data = request.get_json(silent=True) or {}
    titlu = data.get('titlu')
    autor = data.get('autor')
    isbn = data.get('ISBN')
    stoc_total = data.get('stoc_total', 1)
    stoc_disponibil = data.get('stoc_disponibil', stoc_total)
    gen = data.get('gen')
    pozitie = data.get('pozitie') or None
    cod = data.get('cod') or None

    if not titlu or not autor or not isbn or not gen:
        return jsonify({'success': False, 'message': 'titlu, autor, ISBN și gen sunt obligatorii'}), 400

    titlu = str(titlu).strip()
    autor = str(autor).strip()
    isbn = str(isbn).strip().replace('-', '').replace(' ', '')
    gen = str(gen).strip()

    if not titlu or len(titlu) > 50:
        return jsonify({'success': False, 'message': 'Titlul nu poate fi gol și trebuie să aibă maximum 50 de caractere'}), 400
    if not autor or len(autor) > 50:
        return jsonify({'success': False, 'message': 'Autorul nu poate fi gol și trebuie să aibă maximum 50 de caractere'}), 400
    if not isbn or len(isbn) not in (10, 13) or not isbn.isalnum():
        return jsonify({'success': False, 'message': 'ISBN-ul trebuie să aibă 10 sau 13 caractere alfanumerice'}), 400
    if not gen or len(gen) > 255:
        return jsonify({'success': False, 'message': 'Genul nu poate fi gol și trebuie să aibă maximum 255 de caractere'}), 400

    if pozitie:
        pozitie = str(pozitie).strip()
        if len(pozitie) > 100:
            return jsonify({'success': False, 'message': 'Poziția trebuie să aibă maximum 100 de caractere'}), 400
    if cod:
        cod = str(cod).strip()
        if len(cod) > 50:
            return jsonify({'success': False, 'message': 'Codul trebuie să aibă maximum 50 de caractere'}), 400

    try:
        stoc_total = int(stoc_total)
        stoc_disponibil = int(stoc_disponibil)
    except (ValueError, TypeError):
        return jsonify({'success': False, 'message': 'Stocul total și disponibil trebuie să fie numere întregi'}), 400

    if stoc_total < 0 or stoc_disponibil < 0 or stoc_disponibil > stoc_total:
        return jsonify({'success': False, 'message': 'Stocul total/disponibil invalid'}), 400

    insert_query = text(
        "INSERT INTO carti (titlu, autor, ISBN, stoc_total, stoc_disponibil, imprumutat, gen, pozitie, cod) "
        "VALUES (:titlu, :autor, :ISBN, :stoc_total, :stoc_disponibil, :imprumutat, :gen, :pozitie, :cod)"
    )
    try:
        db.session.execute(insert_query, {
            'titlu': titlu,
            'autor': autor,
            'ISBN': isbn,
            'stoc_total': stoc_total,
            'stoc_disponibil': stoc_disponibil,
            'imprumutat': stoc_disponibil < stoc_total,
            'gen': gen,
            'pozitie': pozitie,
            'cod': cod
        })
        db.session.commit()
    except IntegrityError:
        db.session.rollback()
        return jsonify({'success': False, 'message': 'Există deja o carte cu acest ISBN'}), 409
    except Exception:
        db.session.rollback()
        return jsonify({'success': False, 'message': 'Eroare la adăugarea cărții'}), 500

    return jsonify({'success': True, 'message': 'Carte adăugată cu succes'}), 201


@main_bp.route('/books/<int:carte_id>', methods=['PUT'])
@bibliotecar_required
def update_book(carte_id):
    data = request.get_json(silent=True) or {}

    # Construiește query-ul dinamic de update
    fields = {}
    for key in ALLOWED_BOOK_FIELDS:
        if key in data:
            fields[key] = data[key]

    if not fields:
        return jsonify({'success': False, 'message': 'Niciun câmp de actualizat'}), 400

    # Validare input
    if 'titlu' in fields:
        fields['titlu'] = str(fields['titlu']).strip()
        if not fields['titlu'] or len(fields['titlu']) > 50:
            return jsonify({'success': False, 'message': 'Titlul nu poate fi gol și trebuie să aibă maximum 50 de caractere'}), 400
    if 'autor' in fields:
        fields['autor'] = str(fields['autor']).strip()
        if not fields['autor'] or len(fields['autor']) > 50:
            return jsonify({'success': False, 'message': 'Autorul nu poate fi gol și trebuie să aibă maximum 50 de caractere'}), 400
    if 'ISBN' in fields:
        isbn = str(fields['ISBN']).strip().replace('-', '').replace(' ', '')
        if not isbn or len(isbn) not in (10, 13) or not isbn.isalnum():
            return jsonify({'success': False, 'message': 'ISBN-ul trebuie să aibă 10 sau 13 caractere alfanumerice'}), 400
        fields['ISBN'] = isbn
    if 'gen' in fields:
        fields['gen'] = str(fields['gen']).strip()
        if not fields['gen'] or len(fields['gen']) > 255:
            return jsonify({'success': False, 'message': 'Genul nu poate fi gol și trebuie să aibă maximum 255 de caractere'}), 400
    if 'pozitie' in fields:
        if fields['pozitie']:
            fields['pozitie'] = str(fields['pozitie']).strip()
            if len(fields['pozitie']) > 100:
                return jsonify({'success': False, 'message': 'Poziția trebuie să aibă maximum 100 de caractere'}), 400
    if 'cod' in fields:
        if fields['cod']:
            fields['cod'] = str(fields['cod']).strip()
            if len(fields['cod']) > 50:
                return jsonify({'success': False, 'message': 'Codul trebuie să aibă maximum 50 de caractere'}), 400

    if 'stoc_total' in fields or 'stoc_disponibil' in fields:
        book_row = db.session.execute(
            text("SELECT stoc_total, stoc_disponibil FROM carti WHERE carte_id = :id"),
            {'id': carte_id}
        ).fetchone()
        if not book_row:
            return jsonify({'success': False, 'message': 'Cartea nu a fost găsită'}), 404
        
        current_total, current_disponibil = book_row
        new_total = fields.get('stoc_total', current_total)
        new_disponibil = fields.get('stoc_disponibil', current_disponibil)
        
        try:
            new_total = int(new_total)
            new_disponibil = int(new_disponibil)
        except (ValueError, TypeError):
            return jsonify({'success': False, 'message': 'Stocul total și disponibil trebuie să fie numere întregi'}), 400
            
        if new_total < 0 or new_disponibil < 0 or new_disponibil > new_total:
            return jsonify({'success': False, 'message': 'Stocul total/disponibil invalid'}), 400
            
        fields['stoc_total'] = new_total
        fields['stoc_disponibil'] = new_disponibil

    if 'stoc_total' in fields or 'stoc_disponibil' in fields:
        fields['imprumutat'] = fields['stoc_disponibil'] < fields['stoc_total']

    try:
        updated_rows = db.session.query(Carti).filter_by(carte_id=carte_id).update(fields)
        db.session.commit()
        if updated_rows == 0:
            return jsonify({'success': False, 'message': 'Cartea nu a fost găsită'}), 404
    except IntegrityError:
        db.session.rollback()
        return jsonify({'success': False, 'message': 'Conflict ISBN'}), 409
    except Exception:
        db.session.rollback()
        return jsonify({'success': False, 'message': 'Eroare la actualizarea cărții'}), 500

    return jsonify({'success': True, 'message': 'Carte actualizată cu succes'}), 200


@main_bp.route('/books/<int:carte_id>', methods=['DELETE'])
@bibliotecar_required
def delete_book(carte_id):
    try:
        db.session.execute(text("DELETE FROM imprumuturi_active WHERE carte_id = :id"), {'id': carte_id})
        db.session.execute(text("DELETE FROM cereri_carti WHERE carte_id = :id"), {'id': carte_id})
        db.session.execute(text("DELETE FROM recenzii WHERE carte_id = :id"), {'id': carte_id})
        db.session.execute(text("DELETE FROM carti_citite WHERE carte_id = :id"), {'id': carte_id})
        result = db.session.execute(text("DELETE FROM carti WHERE carte_id = :id"), {'id': carte_id})
        db.session.commit()
        if result.rowcount == 0:
            return jsonify({'success': False, 'message': 'Cartea nu a fost găsită'}), 404
    except Exception:
        db.session.rollback()
        return jsonify({'success': False, 'message': 'Eroare la ștergerea cărții'}), 500

    return jsonify({'success': True, 'message': 'Carte ștearsă cu succes'}), 200


@main_bp.route('/books/<int:carte_id>/request-fizic', methods=['POST'])
@jwt_required
def request_book_fizic(carte_id):
    # Notifica toti bibliotecarii pe email
    user = request.current_user

    # Căutăm cartea în baza de date
    book_row = db.session.execute(
        text("SELECT titlu, autor, stoc_disponibil FROM carti WHERE carte_id = :id"),
        {'id': carte_id}
    ).fetchone()

    if not book_row:
        return jsonify({'success': False, 'message': 'Cartea nu a fost găsită'}), 404

    titlu, autor, stoc_disponibil = book_row
    if stoc_disponibil <= 0:
        return jsonify({'success': False, 'message': 'Cartea nu este disponibilă'}), 409

    # Preluăm emailurile tuturor bibliotecarilor
    rows = db.session.execute(
        text("SELECT email FROM users WHERE rol = 'bibliotecar'")
    ).fetchall()
    emailuri_bibliotecari = [r[0] for r in rows if r[0]]

    if not emailuri_bibliotecari:
        logger.warning('Niciun bibliotecar găsit pentru notificarea cererii %d', carte_id)
        return jsonify({'success': False, 'message': 'Niciun bibliotecar disponibil să proceseze cererea'}), 503

    subject = f'Cerere imprumut carte: {titlu}'
    html_body = (
        f'Cerere noua de imprumut\n'
        f'------------------------\n\n'
        f'Utilizator: {user["username"]} ({user["email"]})\n\n'
        f'Carte solicitata:\n'
        f'  Titlu:  {titlu}\n'
        f'  Autor:  {autor}\n'
        f'  ID:     {carte_id}\n\n'
        f'Stoc disponibil: {stoc_disponibil}\n\n'
        f'Intra pe platforma pentru a aproba sau respinge cererea.\n\n'
        f'---\n'
        f'Biblioteca CNI Suceava\n'
        f'Mesaj automat, te rugam nu raspunde la acest email.'
    )

    # Salvăm cererea în baza de date
    try:
        db.session.execute(
            text("INSERT INTO cereri_carti (user_id, carte_id) VALUES (:uid, :cid)"),
            {'uid': user['user_id'], 'cid': carte_id}
        )
        db.session.commit()
    except Exception:
        db.session.rollback()
        logger.exception('Eroare la salvarea cererii în DB')
        return jsonify({'success': False, 'message': 'Eroare la salvarea cererii'}), 500

    # Trimitem emailul (best-effort — cererea nu eșuează dacă emailul nu merge)
    send_email(emailuri_bibliotecari, subject, html_body)

    return jsonify({'success': True, 'message': 'Cererea a fost trimisă!'}), 200


@main_bp.route('/book-requests', methods=['GET'])
@bibliotecar_required
def get_book_requests():
    status_filter = request.args.get('status')  # opțional: pending, approved, rejected
    try:
        sql = (
            "SELECT c.cerere_id, c.user_id, c.carte_id, c.status, c.created_at, "
            "u.username, u.email, ca.titlu, ca.autor "
            "FROM cereri_carti c "
            "JOIN users u ON c.user_id = u.user_id "
            "JOIN carti ca ON c.carte_id = ca.carte_id "
        )
        params = {}
        if status_filter in ('pending', 'approved', 'rejected'):
            sql += "WHERE c.status = :status "
            params['status'] = status_filter
        sql += "ORDER BY c.created_at DESC"

        rows = db.session.execute(text(sql), params).fetchall()
        cereri = []
        for r in rows:
            cereri.append({
                'cerere_id': r[0],
                'user_id': r[1],
                'carte_id': r[2],
                'status': r[3],
                'created_at': r[4].isoformat() if r[4] else None,
                'username': r[5],
                'email': r[6],
                'titlu': r[7],
                'autor': r[8]
            })
        return jsonify({'success': True, 'cereri': cereri}), 200
    except Exception:
        logger.exception('Eroare la preluarea cererilor de împrumut')
        return jsonify({'success': False, 'error': 'Eroare la preluarea cererilor'}), 500


@main_bp.route('/book-requests/<int:cerere_id>', methods=['PUT'])
@bibliotecar_required
def update_book_request(cerere_id):
    data = request.get_json(silent=True) or {}
    status_nou = data.get('status')
    if status_nou not in ('approved', 'rejected'):
        return jsonify({'success': False, 'message': 'Status trebuie să fie approved sau rejected'}), 400

    # Intervalul de ridicare (obligatoriu doar la aprobare)
    ridicare_de_la_str = (data.get('ridicare_de_la') or '').strip()
    ridicare_pana_la_str = (data.get('ridicare_pana_la') or '').strip()
    if status_nou == 'approved' and (not ridicare_de_la_str or not ridicare_pana_la_str):
        return jsonify({'success': False, 'message': 'ridicare_de_la și ridicare_pana_la sunt obligatorii la aprobare'}), 400

    try:
        # Preluăm cererea + datele elevului într-un singur join
        row = db.session.execute(
            text(
                "SELECT cr.user_id, cr.carte_id, cr.status, "
                "u.email, u.username, c.titlu, c.autor "
                "FROM cereri_carti cr "
                "JOIN users u ON cr.user_id = u.user_id "
                "JOIN carti c ON cr.carte_id = c.carte_id "
                "WHERE cr.cerere_id = :id"
            ),
            {'id': cerere_id}
        ).fetchone()
        if not row:
            return jsonify({'success': False, 'message': 'Cererea nu a fost găsită'}), 404

        user_id, carte_id, status_vechi, email_elev, nume_elev, titlu, autor = row

        # Stocăm intervalul de ridicare dacă cererea e aprobată
        ridicare_de_la_dt = None
        ridicare_pana_la_dt = None
        if status_nou == 'approved':
            try:
                ridicare_de_la_dt  = datetime.datetime.fromisoformat(ridicare_de_la_str)
                ridicare_pana_la_dt = datetime.datetime.fromisoformat(ridicare_pana_la_str)
            except ValueError:
                return jsonify({'success': False, 'message': 'Format dată invalid pentru intervalul de ridicare'}), 400

        db.session.execute(
            text("UPDATE cereri_carti SET status = :status, ridicare_de_la = :pf, ridicare_pana_la = :pu WHERE cerere_id = :id"),
            {'status': status_nou, 'pf': ridicare_de_la_dt, 'pu': ridicare_pana_la_dt, 'id': cerere_id}
        )

        # La aprobare trimitem email de notificare (fără să modificăm stocul sau imprumuturi_active)
        if status_nou == 'approved' and status_vechi != 'approved':
            interval_de_la   = ridicare_de_la_dt.strftime('%d.%m.%Y %H:%M')
            interval_pana_la = ridicare_pana_la_dt.strftime('%d.%m.%Y %H:%M')
            html_body = (
                f'Salut, {nume_elev}!\n\n'
                f'Cererea ta de imprumut pentru cartea:\n\n'
                f'  Titlu: {titlu}\n'
                f'  Autor: {autor}\n\n'
                f'a fost APROBATA.\n\n'
                f'Poti ridica cartea de la biblioteca in intervalul:\n'
                f'  {interval_de_la} - {interval_pana_la}\n\n'
                f'Daca nu poti ridica cartea in acest interval, te rugam sa contactezi biblioteca.\n\n'
                f'---\n'
                f'Biblioteca CNI Suceava\n'
                f'Mesaj automat, te rugam nu raspunde la acest email.'
            )
            send_email(
                recipients=[email_elev],
                subject=f'Cerere aprobata: {titlu}',
                body=html_body
            )

        db.session.commit()
        return jsonify({'success': True, 'message': f'Cerere {status_nou}'}), 200
    except Exception:
        db.session.rollback()
        logger.exception('Eroare la actualizarea cererii %d', cerere_id)
        return jsonify({'success': False, 'message': 'Eroare la actualizarea cererii'}), 500


@main_bp.route('/book-requests/<int:cerere_id>/confirma-ridicare', methods=['POST'])
@bibliotecar_required
def confirma_ridicare(cerere_id):
    # Scade stoc, adauga in imprumuturi_active, marcheaza cererea ca ridicat
    try:
        row = db.session.execute(
            text(
                "SELECT cr.user_id, cr.carte_id, cr.status, c.stoc_disponibil "
                "FROM cereri_carti cr "
                "JOIN carti c ON cr.carte_id = c.carte_id "
                "WHERE cr.cerere_id = :id"
            ),
            {'id': cerere_id}
        ).fetchone()

        if not row:
            return jsonify({'success': False, 'message': 'Cererea nu a fost găsită'}), 404

        user_id, carte_id, status_curent, stoc_disponibil = row

        if status_curent != 'approved':
            return jsonify({'success': False, 'message': 'Cererea trebuie să fie aprobată înainte de confirmare'}), 409

        if stoc_disponibil <= 0:
            return jsonify({'success': False, 'message': 'Stoc insuficient — cartea nu mai este disponibilă'}), 409

        now = datetime.datetime.utcnow()
        due = now + datetime.timedelta(days=30)

        # Inserăm în imprumuturi_active
        db.session.execute(
            text("INSERT INTO imprumuturi_active (user_id, carte_id, data_imprumut, data_scadenta) "
                 "VALUES (:uid, :cid, :now, :due)"),
            {'uid': user_id, 'cid': carte_id, 'now': now, 'due': due}
        )

        # Decrementăm stocul disponibil
        db.session.execute(
            text("UPDATE carti SET stoc_disponibil = stoc_disponibil - 1, "
                 "imprumutat = (stoc_disponibil - 1 < stoc_total) "
                 "WHERE carte_id = :cid AND stoc_disponibil > 0"),
            {'cid': carte_id}
        )

        # Marcăm cererea ca ridicată
        db.session.execute(
            text("UPDATE cereri_carti SET status = 'ridicat' WHERE cerere_id = :id"),
            {'id': cerere_id}
        )

        db.session.commit()
        return jsonify({'success': True, 'message': 'Ridicare confirmată'}), 200
    except Exception:
        db.session.rollback()
        logger.exception('Eroare la confirmarea ridicării pentru cererea %d', cerere_id)
        return jsonify({'success': False, 'message': 'Eroare la confirmarea ridicării'}), 500


# raport

@main_bp.route('/librarian/report/docx', methods=['GET'])
@bibliotecar_required
def librarian_report_docx():
    try:
        # Toți utilizatorii obișnuiți (elevi)
        students = db.session.execute(
            text("SELECT user_id, username, email FROM users WHERE rol NOT IN ('1','bibliotecar','administrator') ORDER BY username ASC")
        ).fetchall()

        # Cărți citite per utilizator: {user_id: [(titlu, autor), ...]}
        read_rows = db.session.execute(
            text(
                "SELECT cc.user_id, c.titlu, c.autor "
                "FROM carti_citite cc JOIN carti c ON cc.carte_id = c.carte_id"
            )
        ).fetchall()
        read_map = {}
        for r in read_rows:
            read_map.setdefault(r[0], []).append((r[1], r[2]))

        # Împrumuturi active per utilizator (fizic ridicate)
        borrow_rows = db.session.execute(
            text(
                "SELECT ia.user_id, c.titlu, c.autor, ia.data_imprumut, ia.data_scadenta "
                "FROM imprumuturi_active ia JOIN carti c ON ia.carte_id = c.carte_id "
                "ORDER BY ia.data_scadenta ASC"
            )
        ).fetchall()
        borrow_map = {}
        for r in borrow_rows:
            borrow_map.setdefault(r[0], []).append((r[1], r[2], r[3], r[4]))

        # Cereri aprobate — în așteptarea ridicării fizice
        approved_rows = db.session.execute(
            text(
                "SELECT cr.user_id, c.titlu, c.autor, cr.ridicare_de_la, cr.ridicare_pana_la "
                "FROM cereri_carti cr JOIN carti c ON cr.carte_id = c.carte_id "
                "WHERE cr.status = 'approved' "
                "ORDER BY cr.updated_at ASC"
            )
        ).fetchall()
        approved_map = {}
        for r in approved_rows:
            approved_map.setdefault(r[0], []).append((r[1], r[2], r[3], r[4]))

        # Construim documentul Word
        doc = Document()

        # Titlu
        title_para = doc.add_heading('Raport Bibliotecă – Elevi și Împrumuturi', level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        generated_para = doc.add_paragraph(
            f'Generat la: {datetime.datetime.now().strftime("%d.%m.%Y %H:%M")}'
        )
        generated_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # spațiu gol

        if not students:
            doc.add_paragraph('Nu există elevi înregistrați.')
        else:
            for uid, username, email in students:
                # Heading elev
                doc.add_heading(username, level=1)
                email_para = doc.add_paragraph()
                email_run = email_para.add_run(f'Email: {email}')
                email_run.italic = True

                # Aprobate — în așteptarea ridicării fizice
                approved = approved_map.get(uid, [])
                doc.add_heading('Aprobate — în așteptarea ridicării', level=2)
                if approved:
                    tbl0 = doc.add_table(rows=1, cols=4)
                    tbl0.style = 'Table Grid'
                    hdr0 = tbl0.rows[0].cells
                    hdr0[0].text = 'Titlu'; hdr0[1].text = 'Autor'
                    hdr0[2].text = 'Interval ridicare (de la)'; hdr0[3].text = 'Interval ridicare (până la)'
                    for cell in hdr0:
                        cell.paragraphs[0].runs[0].bold = True
                    for titlu, autor, pf, pu in approved:
                        rc = tbl0.add_row().cells
                        rc[0].text = titlu; rc[1].text = autor
                        rc[2].text = pf.strftime('%d.%m.%Y %H:%M') if pf else '—'
                        rc[3].text = pu.strftime('%d.%m.%Y %H:%M') if pu else '—'
                else:
                    doc.add_paragraph('Nicio carte în așteptarea ridicării.')

                # Cărți împrumutate în prezent
                borrows = borrow_map.get(uid, [])
                doc.add_heading('Cărți împrumutate activ', level=2)
                if borrows:
                    tbl = doc.add_table(rows=1, cols=4)
                    tbl.style = 'Table Grid'
                    hdr = tbl.rows[0].cells
                    hdr[0].text = 'Titlu'
                    hdr[1].text = 'Autor'
                    hdr[2].text = 'Data împrumutului'
                    hdr[3].text = 'Dată scadentă (30 zile)'
                    for cell in hdr:
                        run = cell.paragraphs[0].runs[0]
                        run.bold = True
                    for titlu, autor, data_imp, data_sc in borrows:
                        row_cells = tbl.add_row().cells
                        row_cells[0].text = titlu
                        row_cells[1].text = autor
                        row_cells[2].text = data_imp.strftime('%d.%m.%Y') if data_imp else '—'
                        row_cells[3].text = data_sc.strftime('%d.%m.%Y') if data_sc else '—'
                else:
                    doc.add_paragraph('Niciun împrumut activ.')

                # Cărți deja citite
                reads = read_map.get(uid, [])
                doc.add_heading('Cărți citite', level=2)
                if reads:
                    tbl2 = doc.add_table(rows=1, cols=2)
                    tbl2.style = 'Table Grid'
                    hdr2 = tbl2.rows[0].cells
                    hdr2[0].text = 'Titlu'
                    hdr2[1].text = 'Autor'
                    for cell in hdr2:
                        run = cell.paragraphs[0].runs[0]
                        run.bold = True
                    for titlu, autor in reads:
                        row_cells = tbl2.add_row().cells
                        row_cells[0].text = titlu
                        row_cells[1].text = autor
                else:
                    doc.add_paragraph('Nicio carte citită înregistrată.')

                doc.add_paragraph()  # spațiu între elevi

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        filename = f'raport_biblioteca_{datetime.datetime.now().strftime("%Y%m%d_%H%M")}.docx'
        return send_file(
            buf,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception:
        logger.exception('Eroare la generarea raportului Word')
        return jsonify({'success': False, 'message': 'Eroare la generarea raportului'}), 500


@main_bp.route('/admin/users', methods=['GET'])
@bibliotecar_required
def admin_get_users():
    try:
        rows = db.session.execute(
            text("SELECT user_id, username, email, rol, telefon FROM users ORDER BY username ASC")
        ).fetchall()
        users = [
            {
                'user_id': r[0],
                'username': r[1],
                'email': r[2],
                'rol': normalize_role(r[3]),
                'telefon': r[4]
            }
            for r in rows
        ]
        return jsonify({'success': True, 'users': users}), 200
    except Exception:
        logger.exception('Eroare la preluarea listei de utilizatori')
        return jsonify({'success': False, 'error': 'Eroare la preluarea utilizatorilor'}), 500


@main_bp.route('/admin/users/<int:user_id>', methods=['GET'])
@bibliotecar_required
def admin_get_user_detail(user_id):
    try:
        user_row = db.session.execute(
            text("SELECT user_id, username, email, rol, telefon, description FROM users WHERE user_id = :uid"),
            {'uid': user_id}
        ).fetchone()
        if not user_row:
            return jsonify({'success': False, 'message': 'Utilizatorul nu a fost găsit'}), 404

        # Cărți împrumutate în prezent — sursa autoritativă este imprumuturi_active
        borrowed_rows = db.session.execute(
            text("""
                SELECT ia.imprumut_id, c.carte_id, c.titlu, c.autor, c.ISBN,
                       ia.data_imprumut, ia.data_scadenta
                FROM imprumuturi_active ia
                JOIN carti c ON ia.carte_id = c.carte_id
                WHERE ia.user_id = :uid
                ORDER BY ia.data_imprumut DESC
            """),
            {'uid': user_id}
        ).fetchall()

        # Cărți citite / returnate
        read_rows = db.session.execute(
            text("""
                SELECT c.carte_id, c.titlu, c.autor, c.ISBN
                FROM carti_citite cc
                JOIN carti c ON cc.carte_id = c.carte_id
                WHERE cc.user_id = :uid
            """),
            {'uid': user_id}
        ).fetchall()

        # Istoricul complet al cererilor de împrumut
        history_rows = db.session.execute(
            text("""
                SELECT cc.cerere_id, c.titlu, c.autor, cc.status, cc.created_at,
                       cc.ridicare_de_la, cc.ridicare_pana_la
                FROM cereri_carti cc
                JOIN carti c ON cc.carte_id = c.carte_id
                WHERE cc.user_id = :uid
                ORDER BY cc.created_at DESC
            """),
            {'uid': user_id}
        ).fetchall()

        # Recenzii scrise de utilizator
        review_rows = db.session.execute(
            text("""
                SELECT r.id, c.titlu, c.autor, r.nota, r.comentariu
                FROM recenzii r
                JOIN carti c ON r.carte_id = c.carte_id
                WHERE r.user_id = :uid
                ORDER BY r.id DESC
            """),
            {'uid': user_id}
        ).fetchall()

        return jsonify({
            'success': True,
            'user': {
                'user_id': user_row[0],
                'username': user_row[1],
                'email': user_row[2],
                'rol': normalize_role(user_row[3]),
                'telefon': user_row[4],
                'description': user_row[5]
            },
            'books_borrowed': [
                {
                    'imprumut_id': r[0],
                    'carte_id': r[1], 'titlu': r[2], 'autor': r[3], 'ISBN': r[4],
                    'borrowed_at': r[5].isoformat() if r[5] else None,
                    'due_at': r[6].isoformat() if r[6] else None
                }
                for r in borrowed_rows
            ],
            'books_read': [
                {'carte_id': r[0], 'titlu': r[1], 'autor': r[2], 'ISBN': r[3]}
                for r in read_rows
            ],
            'borrow_history': [
                {
                    'cerere_id': r[0], 'titlu': r[1], 'autor': r[2],
                    'status': r[3],
                    'created_at': r[4].isoformat() if r[4] else None,
                    'ridicare_de_la': r[5].strftime('%d.%m.%Y %H:%M') if r[5] else None,
                    'ridicare_pana_la': r[6].strftime('%d.%m.%Y %H:%M') if r[6] else None
                }
                for r in history_rows
            ],
            'reviews': [
                {'id': r[0], 'titlu': r[1], 'autor': r[2], 'nota': r[3], 'comentariu': r[4]}
                for r in review_rows
            ]
        }), 200
    except Exception:
        logger.exception('Eroare la preluarea detaliilor utilizatorului %d', user_id)
        return jsonify({'success': False, 'error': 'Eroare la preluarea detaliilor utilizatorului'}), 500


@main_bp.route('/admin/loans/<int:imprumut_id>/prelungeste', methods=['PUT'])
@bibliotecar_required
def prelungeste_imprumut(imprumut_id):
    data = request.get_json(silent=True) or {}
    data_noua = (data.get('data_scadenta') or '').strip()
    if not data_noua:
        return jsonify({'success': False, 'message': 'data_scadenta este obligatorie'}), 400
    try:
        data_scadenta = datetime.datetime.fromisoformat(data_noua)
    except ValueError:
        return jsonify({'success': False, 'message': 'Format dată invalid. Folosește YYYY-MM-DDTHH:MM'}), 400
    try:
        result = db.session.execute(
            text("UPDATE imprumuturi_active SET data_scadenta = :ds WHERE imprumut_id = :id"),
            {'ds': data_scadenta, 'id': imprumut_id}
        )
        if result.rowcount == 0:
            return jsonify({'success': False, 'message': 'Împrumutul nu a fost găsit'}), 404
        db.session.commit()
        return jsonify({'success': True, 'message': 'Data scadentă actualizată'}), 200
    except Exception:
        db.session.rollback()
        logger.exception('Eroare la prelungire împrumut %d', imprumut_id)
        return jsonify({'success': False, 'message': 'Eroare la actualizare'}), 500


@main_bp.route('/admin/loans/<int:imprumut_id>/returneaza', methods=['POST'])
@bibliotecar_required
def returneaza_imprumut(imprumut_id):
    # Sterge din imprumuturi_active, creste stocul, adauga in carti_citite
    try:
        row = db.session.execute(
            text("SELECT user_id, carte_id FROM imprumuturi_active WHERE imprumut_id = :id"),
            {'id': imprumut_id}
        ).fetchone()
        if not row:
            return jsonify({'success': False, 'message': 'Împrumutul nu a fost găsit'}), 404
        user_id, carte_id = row

        db.session.execute(
            text("DELETE FROM imprumuturi_active WHERE imprumut_id = :id"),
            {'id': imprumut_id}
        )
        db.session.execute(
            text("UPDATE carti SET stoc_disponibil = stoc_disponibil + 1, "
                 "imprumutat = (stoc_disponibil + 1 < stoc_total) "
                 "WHERE carte_id = :cid"),
            {'cid': carte_id}
        )

        already = db.session.execute(
            text("SELECT 1 FROM carti_citite WHERE user_id = :uid AND carte_id = :cid"),
            {'uid': user_id, 'cid': carte_id}
        ).fetchone()
        if not already:
            db.session.execute(
                text("INSERT INTO carti_citite (user_id, carte_id) VALUES (:uid, :cid)"),
                {'uid': user_id, 'cid': carte_id}
            )

        db.session.commit()
        return jsonify({'success': True, 'message': 'Carte returnată cu succes'}), 200
    except Exception:
        db.session.rollback()
        logger.exception('Eroare la returnarea împrumutului %d', imprumut_id)
        return jsonify({'success': False, 'message': 'Eroare la returnare'}), 500


@main_bp.route('/reviews', methods=['GET'])
def get_reviews():
    carte_id = request.args.get('carte_id')
    if not carte_id:
        return jsonify({'success': False, 'message': 'carte_id este obligatoriu'}), 400

    try:
        query = text("""
            SELECT r.id, r.nota, r.comentariu, u.username
            FROM recenzii r
            JOIN users u ON r.user_id = u.user_id
            WHERE r.carte_id = :carte_id
            ORDER BY r.id DESC
        """)
        result = db.session.execute(query, {'carte_id': carte_id})
        reviews = []
        for row in result:
            reviews.append({
                'id': row[0],
                'nota': row[1],
                'comentariu': row[2],
                'username': row[3]
            })

        # Calculăm media notelor
        avg_rating = 0
        if reviews:
            avg_rating = round(sum(r['nota'] for r in reviews) / len(reviews), 1)

        return jsonify({
            'success': True,
            'reviews': reviews,
            'avg_rating': avg_rating,
            'total_reviews': len(reviews)
        }), 200
    except Exception:
        logger.exception('Eroare la preluarea recenziilor')
        return jsonify({'success': False, 'error': 'Eroare la preluarea recenziilor'}), 500


@main_bp.route('/reviews', methods=['POST'])
@jwt_required
def submit_review():
    user = request.current_user
    data = request.get_json(silent=True) or {}
    carte_id = data.get('carte_id')
    nota = data.get('nota')
    comentariu = (data.get('comentariu') or '').strip()

    if not carte_id or nota is None or not comentariu:
        return jsonify({'success': False, 'message': 'carte_id, nota și comentariu sunt obligatorii'}), 400

    try:
        nota = int(nota)
    except (ValueError, TypeError):
        return jsonify({'success': False, 'message': 'nota trebuie să fie un număr întreg 1-5'}), 400

    if nota < 1 or nota > 5:
        return jsonify({'success': False, 'message': 'nota trebuie să fie între 1 și 5'}), 400

    # Verificăm că cartea există
    book = db.session.execute(text("SELECT carte_id FROM carti WHERE carte_id = :id"), {'id': carte_id}).fetchone()
    if not book:
        return jsonify({'success': False, 'message': 'Cartea nu a fost găsită'}), 404

    try:
        # Upsert: o recenzie per utilizator per carte
        existing = db.session.execute(
            text("SELECT id FROM recenzii WHERE user_id = :uid AND carte_id = :cid"),
            {'uid': user['user_id'], 'cid': carte_id}
        ).fetchone()

        if existing:
            db.session.execute(
                text("UPDATE recenzii SET nota = :nota, comentariu = :comentariu WHERE id = :id"),
                {'nota': nota, 'comentariu': comentariu, 'id': existing[0]}
            )
        else:
            db.session.execute(
                text("INSERT INTO recenzii (user_id, carte_id, nota, comentariu) VALUES (:uid, :cid, :nota, :com)"),
                {'uid': user['user_id'], 'cid': carte_id, 'nota': nota, 'com': comentariu}
            )
        db.session.commit()
        return jsonify({'success': True, 'message': 'Recenzie salvată'}), 200
    except Exception:
        db.session.rollback()
        logger.exception('Eroare la salvarea recenziei')
        return jsonify({'success': False, 'message': 'Eroare la salvarea recenziei'}), 500


@main_bp.route('/reviews/user', methods=['GET'])
@jwt_required
def get_user_reviews():
    user_id = request.current_user['user_id']

    try:
        query = text("""
            SELECT r.id, r.nota, r.comentariu, c.titlu, c.autor
            FROM recenzii r
            JOIN carti c ON r.carte_id = c.carte_id
            WHERE r.user_id = :user_id
            ORDER BY r.id DESC
        """)
        result = db.session.execute(query, {'user_id': user_id})
        reviews = [
            {'id': row[0], 'nota': row[1], 'comentariu': row[2], 'titlu': row[3], 'autor': row[4]}
            for row in result
        ]
        return jsonify({'success': True, 'reviews': reviews, 'total_reviews': len(reviews)}), 200
    except Exception:
        logger.exception('Eroare la preluarea recenziilor utilizatorului')
        return jsonify({'success': False, 'error': 'Eroare la preluarea recenziilor'}), 500


# autentificare
@main_bp.route('/auth/login', methods=['POST'])
def login():
    data = request.get_json(silent=True) or {}
    raw_email = data.get('email')
    parola = data.get('password')

    if not raw_email or not parola:
        return jsonify({'success': False, 'message': 'Email și parola sunt obligatorii'}), 400

    email = normalize_cni_email(raw_email)
    if not email:
        return jsonify({'success': False, 'message': 'Doar emailurile @cni-sv.ro sunt permise'}), 400

    result = db.session.execute(
        text("SELECT user_id, username, email, hashed_password, rol FROM users WHERE email = :email"),
        {'email': email}
    ).mappings().first()

    if not result:
        return jsonify({'success': False, 'message': 'Credențiale invalide'}), 401

    parola_hash = result.get('hashed_password')
    if not parola_hash or not bcrypt.checkpw(parola.encode('utf-8'), parola_hash.encode('utf-8')):
        return jsonify({'success': False, 'message': 'Credențiale invalide'}), 401

    # Șterge orice cod vechi pentru acest utilizator
    db.session.execute(
        text("DELETE FROM coduri_verificare WHERE user_id = :uid"),
        {'uid': result['user_id']}
    )

    # Generează cod 6 cifre și token temporar
    code = f"{random.randint(0, 999999):06d}"
    temp_token = secrets.token_urlsafe(32)
    expires_at = datetime.datetime.utcnow() + datetime.timedelta(minutes=10)

    db.session.execute(
        text("INSERT INTO coduri_verificare (user_id, code, temp_token, expires_at) VALUES (:uid, :code, :token, :exp)"),
        {'uid': result['user_id'], 'code': code, 'token': temp_token, 'exp': expires_at}
    )
    db.session.commit()

    # Trimite codul pe email
    body = (
        f'Salut, {result["username"]}!\n\n'
        f'Codul tau de verificare pentru autentificarea in Biblioteca CNI este:\n\n'
        f'    {code}\n\n'
        f'Codul este valabil 10 minute. Nu il impartasi cu nimeni.\n\n'
        f'Daca nu ai incercat sa te autentifici, ignora acest mesaj.\n\n'
        f'---\n'
        f'Biblioteca CNI Suceava\n'
        f'Mesaj automat, te rugam nu raspunde la acest email.'
    )
    send_email([result['email']], 'Cod de verificare - Biblioteca CNI', body)

    return jsonify({'success': True, 'step': 'verify', 'temp_token': temp_token}), 200


@main_bp.route('/auth/verify-code', methods=['POST'])
def verify_code():
    data = request.get_json(silent=True) or {}
    temp_token = (data.get('temp_token') or '').strip()
    code = (data.get('code') or '').strip()

    if not temp_token or not code:
        return jsonify({'success': False, 'message': 'Token și cod sunt obligatorii'}), 400

    row = db.session.execute(
        text("SELECT vc.id, vc.user_id, vc.code, vc.expires_at, u.username, u.email, u.rol "
             "FROM coduri_verificare vc "
             "JOIN users u ON u.user_id = vc.user_id "
             "WHERE vc.temp_token = :token"),
        {'token': temp_token}
    ).mappings().first()

    if not row:
        return jsonify({'success': False, 'message': 'Sesiune de verificare invalidă'}), 401

    if datetime.datetime.utcnow() > row['expires_at']:
        db.session.execute(text("DELETE FROM coduri_verificare WHERE id = :id"), {'id': row['id']})
        db.session.commit()
        return jsonify({'success': False, 'message': 'Codul a expirat. Încearcă să te autentifici din nou.'}), 401

    if code != row['code']:
        db.session.execute(text("DELETE FROM coduri_verificare WHERE id = :id"), {'id': row['id']})
        db.session.commit()
        return jsonify({'success': False, 'message': 'Cod incorect. Sesiunea de verificare a fost anulată. Te rugăm să încerci din nou.'}), 401

    # Cod valid — șterge înregistrarea și emite JWT
    db.session.execute(text("DELETE FROM coduri_verificare WHERE id = :id"), {'id': row['id']})
    db.session.commit()

    token = create_jwt_token(user_id=row['user_id'], username=row['username'], email=row['email'])
    resp = jsonify({
        'success': True,
        'message': 'Autentificare reușită',
        'user': {
            'user_id': row['user_id'],
            'username': row['username'],
            'email': row['email'],
            'rol': normalize_role(row['rol'])
        }
    })
    set_jwt_cookie(resp, token)
    return resp, 200

@main_bp.route('/auth/me', methods=['GET'])
@jwt_required
def auth_me():
    user = request.current_user
    # Citim câmpul club direct din DB (nu e stocat în JWT)
    row = db.session.execute(
        text("SELECT club FROM users WHERE user_id = :uid"),
        {'uid': user['user_id']}
    ).fetchone()
    return jsonify({
        'success': True,
        'user_id': user['user_id'],
        'username': user['username'],
        'email': user['email'],
        'rol': user['rol'],
        'club': bool(row[0]) if row else False
    }), 200


@main_bp.route('/auth/logout', methods=['POST'])
def logout():
    resp = jsonify({'success': True, 'message': 'Deconectat cu succes'})
    clear_jwt_cookie(resp)
    return resp, 200


@main_bp.route('/auth/forgot-password', methods=['POST'])
def forgot_password():
    data = request.get_json(silent=True) or {}
    email = (data.get('email') or '').strip().lower()

    if not email:
        return jsonify({'success': False, 'message': 'Email-ul este obligatoriu'}), 400

    row = db.session.execute(
        text("SELECT user_id, username FROM users WHERE email = :email"),
        {'email': email}
    ).fetchone()

    if not row:
        # Prevent email enumeration by returning a generic success message
        return jsonify({'success': True, 'message': 'Dacă email-ul există, s-a trimis un link de resetare.'}), 200

    user_id = row[0]
    username = row[1]
    token = secrets.token_urlsafe(32)
    expires_at = datetime.datetime.utcnow() + datetime.timedelta(hours=1)

    try:
        db.session.execute(
            text("INSERT INTO password_reset_tokens (token, user_id, expires_at) VALUES (:token, :uid, :exp)"),
            {'token': token, 'uid': user_id, 'exp': expires_at}
        )
        db.session.commit()
    except Exception:
        db.session.rollback()
        return jsonify({'success': False, 'message': 'Eroare la generarea token-ului.'}), 500

    # Determinăm dinamic adresa de bază a frontend-ului
    frontend_url = request.headers.get('Origin')
    if not frontend_url:
        referer = request.headers.get('Referer')
        if referer:
            # Păstrăm doar origin-ul din referer (ex. https://site.ro/forgot-password -> https://site.ro)
            from urllib.parse import urlparse
            parsed_referer = urlparse(referer)
            frontend_url = f"{parsed_referer.scheme}://{parsed_referer.netloc}"
        else:
            # Fallback dacă antetele lipsesc
            frontend_url = "http://localhost:5173"

    reset_url = f"{frontend_url}/reset-password/{token}"
    
    html_body = (
        f"Salut {username},\n\n"
        f"Am primit o cerere de resetare a parolei pentru contul tău.\n"
        f"Accesează acest link pentru a o reseta:\n"
        f"{reset_url}\n\n"
        f"Acest link este valabil o oră. Dacă nu ai cerut asta, ignoră acest email.\n\n"
        f"Echipa Biblioteca"
    )

    send_email([email], "Resetare Parolă", html_body)

    return jsonify({'success': True, 'message': 'Dacă email-ul există, s-a trimis un link de resetare.'}), 200


@main_bp.route('/auth/reset-password/<token>', methods=['POST'])
def reset_password(token):
    data = request.get_json(silent=True) or {}
    new_password = data.get('password')

    if not new_password or len(new_password) < 6:
        return jsonify({'success': False, 'message': 'Parola trebuie să aibă minim 6 caractere'}), 400

    now = datetime.datetime.utcnow()

    # Preluăm tokenul
    row = db.session.execute(
        text("SELECT user_id, expires_at FROM password_reset_tokens WHERE token = :token"),
        {'token': token}
    ).fetchone()

    if not row:
        return jsonify({'success': False, 'message': 'Token invalid sau expirat.'}), 400
        
    user_id = row[0]
    expires_at = row[1]

    if expires_at < now:
        db.session.execute(text("DELETE FROM password_reset_tokens WHERE token = :token"), {'token': token})
        db.session.commit()
        return jsonify({'success': False, 'message': 'Token expirat.'}), 400

    hashed_pw = bcrypt.hashpw(new_password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

    try:
        db.session.execute(
            text("UPDATE users SET hashed_password = :hpw WHERE user_id = :uid"),
            {'hpw': hashed_pw, 'uid': user_id}
        )
        db.session.execute(
            text("DELETE FROM password_reset_tokens WHERE token = :token"),
            {'token': token}
        )
        db.session.commit()
    except Exception:
        db.session.rollback()
        return jsonify({'success': False, 'message': 'Eroare la resetarea parolei.'}), 500

    return jsonify({'success': True, 'message': 'Parola a fost resetată cu succes!'}), 200


@main_bp.route('/auth/profile', methods=['GET'])
@jwt_required
def profile():
    user = request.current_user

    query = text(
        "SELECT user_id, username, email, description FROM users WHERE user_id = :user_id"
    )
    result = db.session.execute(query, {'user_id': user['user_id']}).mappings().first()
    if not result:
        return jsonify({'success': False, 'message': 'Profilul nu a fost găsit'}), 404

    return jsonify({
        'success': True,
        'user_id': result.get('user_id'),
        'username': result.get('username'),
        'email': result.get('email'),
        'description': result.get('description')
    }), 200

@main_bp.route('/auth/books-read', methods=['GET'])
@jwt_required
def books_read():
    user = request.current_user
    user_id = user['user_id']

    query = text(
        "SELECT c.carte_id, c.titlu, c.autor, c.ISBN "
        "FROM carti_citite cc "
        "JOIN carti c ON cc.carte_id = c.carte_id "
        "WHERE cc.user_id = :user_id"
    )
    rows = db.session.execute(query, {'user_id': user_id}).mappings().all()

    books = []
    for row in rows:
        books.append({
            'carte_id': row.get('carte_id'),
            'titlu': row.get('titlu'),
            'autor': row.get('autor'),
            'ISBN': row.get('ISBN')
        })

    return jsonify({'success': True, 'books': books}), 200

@main_bp.route('/auth/profile', methods=['PUT'])
@jwt_required
def update_profile():
    user = request.current_user
    data = request.get_json(silent=True) or {}
    description = data.get('description')

    if description is not None:
        description = str(description).strip()
        if len(description) > 255:
            return jsonify({'success': False, 'message': 'Descrierea este prea lungă (max 255 caractere)'}), 400

    update_query = text(
        "UPDATE users SET description = :description WHERE user_id = :user_id"
    )
    try:
        result = db.session.execute(update_query, {'description': description, 'user_id': user['user_id']})
        db.session.commit()
        if result.rowcount == 0:
            return jsonify({'success': False, 'message': 'Profilul nu a fost găsit'}), 404
    except Exception:
        db.session.rollback()
        return jsonify({'success': False, 'message': 'Eroare la actualizarea descrierii'}), 500

    return jsonify({'success': True, 'message': 'Descriere actualizată'}), 200

@main_bp.route('/auth/register', methods=['POST'])
def register():
    data = request.get_json(silent=True) or {}
    username = data.get('user') or data.get('username') or data.get('fullName')
    raw_email = data.get('email')
    parola = data.get('password')

    if not username or not raw_email or not parola:
        return jsonify({'success': False, 'message': 'Username, email și parola sunt obligatorii'}), 400

    username = username.strip()
    if not USERNAME_REGEX.match(username):
        return jsonify({'success': False, 'message': 'Numele complet trebuie să aibă între 3 și 50 de caractere și să conțină doar litere, cifre, spații, puncte, cratime sau caractere de subliniere'}), 400

    email = normalize_cni_email(raw_email)
    if not email:
        return jsonify({'success': False, 'message': 'Doar emailurile @cni-sv.ro sunt permise'}), 400

    if len(parola) < 8 or len(parola) > 128:
        return jsonify({'success': False, 'message': 'Parola trebuie să aibă între 8 și 128 de caractere'}), 400

    parola_hash = bcrypt.hashpw(parola.encode('utf-8'), bcrypt.gensalt(12)).decode('utf-8')
    values = {
        'username': username,
        'email': email,
        'hashed_password': parola_hash,
        'rol': 'user',
        'telefon': None,
        'club': 0
    }

    insert_query = text(
        "INSERT INTO users (username, email, hashed_password, rol, telefon, club) VALUES (:username, :email, :hashed_password, :rol, :telefon, :club)"
    )

    try:
        db.session.execute(insert_query, values)
        db.session.commit()
    except IntegrityError:
        db.session.rollback()
        return jsonify({'success': False, 'message': 'Username sau email deja existente'}), 409
    except Exception:
        db.session.rollback()
        return jsonify({'success': False, 'message': 'Înregistrarea a eșuat'}), 500

    return jsonify({'success': True, 'message': 'Cont creat cu succes. Te poti autentifica acum.'}), 201


@main_bp.route('/auth/profile-picture', methods=['POST'])
@jwt_required
def upload_profile_picture():
    user = request.current_user

    if 'file' not in request.files:
        return jsonify({'success': False, 'message': 'Niciun fișier trimis'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': 'Niciun fișier selectat'}), 400

    if not allowed_file(file.filename):
        return jsonify({'success': False, 'message': 'Tip de fișier nepermis. Folosește png, jpg, jpeg, gif sau webp'}), 400

    username = user['username']
    ext = file.filename.rsplit('.', 1)[1].lower()

    # Ștergem poza de profil existentă pentru acest utilizator
    for poza_veche in _find_files_by_prefix(PROFILE_PICTURES_DIR, username):
        os.remove(os.path.join(PROFILE_PICTURES_DIR, poza_veche))

    # Salvăm noul fișier ca username.ext
    filename = f"{username}.{ext}"
    filepath = os.path.join(PROFILE_PICTURES_DIR, filename)
    file.save(filepath)

    return jsonify({
        'success': True,
        'message': 'Poză de profil încărcată',
        'filename': filename
    }), 200


@main_bp.route('/auth/profile-picture/<username>', methods=['GET'])
def get_profile_picture(username):
    # daca nu exista poza, returneaza default.jpg
    matches = _find_files_by_prefix(PROFILE_PICTURES_DIR, username)

    if not matches:
        return send_from_directory(PROFILE_PICTURES_DIR, 'default.jpg')

    return send_from_directory(PROFILE_PICTURES_DIR, matches[0])


@main_bp.route('/books/image', methods=['POST'])
@bibliotecar_required
def upload_book_image():
    # salvat ca <carte_id>.<ext> in book_images/
    carte_id = request.form.get('carte_id')
    if not carte_id:
        return jsonify({'success': False, 'message': 'carte_id is required'}), 400

    if 'file' not in request.files:
        return jsonify({'success': False, 'message': 'Niciun fișier trimis'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': 'Niciun fișier selectat'}), 400

    if not allowed_file(file.filename):
        return jsonify({'success': False, 'message': 'Tip de fișier nepermis. Folosește png, jpg, jpeg, gif sau webp'}), 400

    # Verificăm că cartea există
    query = text("SELECT carte_id FROM carti WHERE carte_id = :carte_id")
    result = db.session.execute(query, {'carte_id': carte_id}).mappings().first()
    if not result:
        return jsonify({'success': False, 'message': 'Cartea nu a fost găsită'}), 404

    ext = file.filename.rsplit('.', 1)[1].lower()

    # Ștergem coperta existentă a cărții
    for imagine_veche in _find_files_by_prefix(BOOK_IMAGES_DIR, str(carte_id)):
        os.remove(os.path.join(BOOK_IMAGES_DIR, imagine_veche))

    # Salvăm noul fișier ca carte_id.ext
    filename = f"{carte_id}.{ext}"
    filepath = os.path.join(BOOK_IMAGES_DIR, filename)
    file.save(filepath)

    return jsonify({
        'success': True,
        'message': 'Copertă carte încărcată',
        'filename': filename
    }), 200


@main_bp.route('/books/image/<int:carte_id>', methods=['GET'])
def get_book_image(carte_id):
    matches = _find_files_by_prefix(BOOK_IMAGES_DIR, str(carte_id))

    if not matches:
        return jsonify({'success': False, 'message': 'Nicio imagine găsită'}), 404

    return send_from_directory(BOOK_IMAGES_DIR, matches[0])


@main_bp.route('/books/pdf', methods=['POST'])
@bibliotecar_required
def upload_book_pdf():
    carte_id = request.form.get('carte_id')
    if not carte_id:
        return jsonify({'success': False, 'message': 'carte_id is required'}), 400
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': 'Niciun fișier trimis'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': 'Niciun fișier selectat'}), 400
    if not file.filename.lower().endswith('.pdf'):
        return jsonify({'success': False, 'message': 'Doar fișiere PDF sunt acceptate'}), 400

    result = db.session.execute(text("SELECT carte_id FROM carti WHERE carte_id = :id"), {'id': carte_id}).mappings().first()
    if not result:
        return jsonify({'success': False, 'message': 'Cartea nu a fost găsită'}), 404

    os.makedirs(BOOK_PDFS_DIR, exist_ok=True)
    old_path = os.path.join(BOOK_PDFS_DIR, f"{carte_id}.pdf")
    if os.path.exists(old_path):
        os.remove(old_path)

    file.save(os.path.join(BOOK_PDFS_DIR, f"{carte_id}.pdf"))
    return jsonify({'success': True, 'message': 'PDF încărcat'}), 200


@main_bp.route('/books/pdf/<int:carte_id>', methods=['GET'])
@jwt_required
def get_book_pdf(carte_id):
    pdf_path = os.path.join(BOOK_PDFS_DIR, f"{carte_id}.pdf")
    if not os.path.exists(pdf_path):
        return jsonify({'success': False, 'message': 'PDF negăsit'}), 404
    return send_from_directory(BOOK_PDFS_DIR, f"{carte_id}.pdf", as_attachment=False)


@main_bp.route('/books/pdf/<int:carte_id>', methods=['DELETE'])
@bibliotecar_required
def delete_book_pdf(carte_id):
    pdf_path = os.path.join(BOOK_PDFS_DIR, f"{carte_id}.pdf")
    if not os.path.exists(pdf_path):
        return jsonify({'success': False, 'message': 'PDF negăsit'}), 404
    os.remove(pdf_path)
    return jsonify({'success': True, 'message': 'PDF șters'}), 200


# anunturi

@main_bp.route('/anunturi', methods=['GET'])
def get_anunturi():
    user = get_current_user()
    user_id = user['user_id'] if user else None

    try:
        query = text("""
            SELECT a.anunt_id, a.titlu, a.anunt, a.data_publicare, a.aprecieri
            FROM anunturi a
            ORDER BY a.data_publicare DESC
        """)
        rows = db.session.execute(query).fetchall()

        liked_set = set()
        if user_id:
            liked_rows = db.session.execute(
                text("SELECT anunt_id FROM anunturi_aprecieri WHERE user_id = :uid"),
                {'uid': user_id}
            ).fetchall()
            liked_set = {r[0] for r in liked_rows}

        anunturi = []
        for row in rows:
            anunturi.append({
                'anunt_id': row[0],
                'titlu': row[1],
                'anunt': row[2],
                'data_publicare': row[3].strftime('%d %B %Y, %H:%M') if row[3] else None,
                'aprecieri': row[4] or 0,
                'liked': row[0] in liked_set
            })

        return jsonify({'success': True, 'anunturi': anunturi}), 200
    except Exception:
        logger.exception('Eroare la preluarea anunțurilor')
        return jsonify({'success': False, 'error': 'Eroare la preluarea anunțurilor'}), 500


@main_bp.route('/anunturi', methods=['POST'])
@bibliotecar_required
def create_anunt():
    data = request.get_json(silent=True) or {}
    titlu = data.get('titlu', '').strip()
    anunt = data.get('anunt', '').strip()

    if not titlu or not anunt:
        return jsonify({'success': False, 'message': 'titlu și anunt sunt obligatorii'}), 400

    try:
        db.session.execute(
            text("INSERT INTO anunturi (titlu, anunt, data_publicare, aprecieri) VALUES (:titlu, :anunt, NOW(), 0)"),
            {'titlu': titlu, 'anunt': anunt}
        )
        db.session.commit()
        return jsonify({'success': True, 'message': 'Anunț creat'}), 201
    except Exception:
        db.session.rollback()
        logger.exception('Eroare la crearea anunțului')
        return jsonify({'success': False, 'message': 'Eroare la crearea anunțului'}), 500


@main_bp.route('/anunturi/<int:anunt_id>', methods=['PUT'])
@bibliotecar_required
def update_anunt(anunt_id):
    data = request.get_json(silent=True) or {}

    fields = {}
    for key in ALLOWED_ANUNT_FIELDS:
        if key in data and data[key] is not None:
            fields[key] = data[key].strip()

    if not fields:
        return jsonify({'success': False, 'message': 'Niciun câmp de actualizat'}), 400

    try:
        updated_rows = db.session.query(Anunturi).filter_by(anunt_id=anunt_id).update(fields)
        db.session.commit()
        if updated_rows == 0:
            return jsonify({'success': False, 'message': 'Anunțul nu a fost găsit'}), 404
        return jsonify({'success': True, 'message': 'Anunț actualizat'}), 200
    except Exception:
        db.session.rollback()
        logger.exception('Eroare la actualizarea anunțului %d', anunt_id)
        return jsonify({'success': False, 'message': 'Eroare la actualizarea anunțului'}), 500


@main_bp.route('/anunturi/<int:anunt_id>', methods=['DELETE'])
@bibliotecar_required
def delete_anunt(anunt_id):
    try:
        result = db.session.execute(text("DELETE FROM anunturi WHERE anunt_id = :id"), {'id': anunt_id})
        db.session.commit()
        if result.rowcount == 0:
            return jsonify({'success': False, 'message': 'Announcement not found'}), 404
        return jsonify({'success': True, 'message': 'Anunț șters'}), 200
    except Exception:
        db.session.rollback()
        logger.exception('Eroare la ștergerea anunțului %d', anunt_id)
        return jsonify({'success': False, 'message': 'Eroare la ștergerea anunțului'}), 500


@main_bp.route('/anunturi/<int:anunt_id>/like', methods=['POST'])
@jwt_required
def toggle_like_anunt(anunt_id):
    user = request.current_user
    user_id = user['user_id']

    try:
        existing = db.session.execute(
            text("SELECT id FROM anunturi_aprecieri WHERE anunt_id = :aid AND user_id = :uid"),
            {'aid': anunt_id, 'uid': user_id}
        ).fetchone()

        if existing:
            db.session.execute(
                text("DELETE FROM anunturi_aprecieri WHERE anunt_id = :aid AND user_id = :uid"),
                {'aid': anunt_id, 'uid': user_id}
            )
            db.session.execute(
                text("UPDATE anunturi SET aprecieri = GREATEST(aprecieri - 1, 0) WHERE anunt_id = :aid"),
                {'aid': anunt_id}
            )
            db.session.commit()
            apreciat = False
        else:
            db.session.execute(
                text("INSERT INTO anunturi_aprecieri (anunt_id, user_id) VALUES (:aid, :uid)"),
                {'aid': anunt_id, 'uid': user_id}
            )
            db.session.execute(
                text("UPDATE anunturi SET aprecieri = aprecieri + 1 WHERE anunt_id = :aid"),
                {'aid': anunt_id}
            )
            db.session.commit()
            apreciat = True

        count = db.session.execute(
            text("SELECT aprecieri FROM anunturi WHERE anunt_id = :aid"),
            {'aid': anunt_id}
        ).fetchone()

        return jsonify({'success': True, 'liked': apreciat, 'aprecieri': count[0] if count else 0}), 200
    except Exception:
        db.session.rollback()
        logger.exception('Eroare la toggle apreciere anunț %d', anunt_id)
        return jsonify({'success': False, 'message': 'Eroare la toggle apreciere'}), 500


# ai

GROQ_MODEL = 'llama-3.3-70b-versatile'

def _get_groq_client():
    api_key = current_app.config.get('GROQ_API_KEY', '')
    if not api_key:
        return None
    return Groq(api_key=api_key)

def _groq_chat(client, prompt, system_message=None):
    messages = []
    if system_message:
        messages.append({'role': 'system', 'content': system_message})
    messages.append({'role': 'user', 'content': prompt})
    
    logger.info('Prompt Groq: %s', prompt[:300])
    raspuns = client.chat.completions.create(
        model=GROQ_MODEL,
        messages=messages,
        max_tokens=1024,
        temperature=0.7
    )
    return raspuns.choices[0].message.content.strip()


@main_bp.route('/ai/recommend', methods=['GET'])
@jwt_required
def ai_recommend():
    user = request.current_user
    user_id = user['user_id']

    client = _get_groq_client()
    if not client:
        return jsonify({'success': False, 'message': 'AI neconfigurat'}), 503

    try:
        # Toate cărțile din bibliotecă cu disponibilitate reală
        all_rows = db.session.execute(
            text("SELECT carte_id, titlu, autor, gen, stoc_disponibil FROM carti ORDER BY titlu")
        ).fetchall()

        # Cărțile pe care utilizatorul le-a citit/returnat deja
        read_rows = db.session.execute(
            text("SELECT c.carte_id, c.titlu, c.autor, c.gen FROM carti_citite cc JOIN carti c ON cc.carte_id = c.carte_id WHERE cc.user_id = :uid"),
            {'uid': user_id}
        ).fetchall()

        # Cărțile pe care utilizatorul le are împrumutate sau solicitate
        borrowed_rows = db.session.execute(
            text("SELECT c.carte_id, c.titlu, c.autor FROM cereri_carti cc JOIN carti c ON cc.carte_id = c.carte_id WHERE cc.user_id = :uid AND cc.status IN ('pending', 'approved')"),
            {'uid': user_id}
        ).fetchall()

        if not read_rows:
            return jsonify({'success': True, 'recommendations': '', 'no_history': True}), 200

        read_list = '\n'.join(f'- {r[1]} de {r[2]} (gen: {r[3]})' for r in read_rows)
        borrowed_list = '\n'.join(f'- {r[1]} de {r[2]}' for r in borrowed_rows) or 'niciuna'
        lib_list = '\n'.join(
            f'- [ID:{r[0]}] {r[1]} de {r[2]} (gen: {r[3]}) — {"disponibil" if r[4] > 0 else "indisponibil"}'
            for r in all_rows
        )

        system_prompt = (
            "Ești bibliotecar la o școală. Ai acces la catalogul COMPLET al bibliotecii de mai jos.\n"
            "Recomandă DOAR cărți care există în catalog și sunt marcate ca 'disponibil'.\n"
            "Nu inventa cărți sau autori care nu sunt în catalog.\n\n"
            f"CATALOG COMPLET AL BIBLIOTECII:\n{lib_list}"
        )
        user_prompt = (
            f"Cărți citite de elev (nu le recomanda din nou):\n{read_list}\n\n"
            f"Cărți deja împrumutate/solicitate de elev:\n{borrowed_list}\n\n"
            "Recomandă 3-4 cărți disponibile din catalog pe care elevul nu le-a citit și nu le are deja. "
            "Pentru fiecare menționează titlul exact din catalog, autorul și de ce i-ar plăcea elevului (1-2 propoziții). "
            "Răspunde în română, prietenos."
        )

        raspuns_ai = _groq_chat(client, user_prompt, system_message=system_prompt)
        return jsonify({'success': True, 'recommendations': raspuns_ai}), 200
    except Exception:
        logger.exception('Eroare AI recommend')
        return jsonify({'success': False, 'message': 'Cererea AI a eșuat'}), 500


@main_bp.route('/ai/review-assist', methods=['POST'])
@jwt_required
def ai_review_assist():
    data = request.get_json(silent=True) or {}
    ciorna = (data.get('draft') or '').strip()
    titlu_carte = (data.get('titlu') or '').strip()

    if not ciorna:
        return jsonify({'success': False, 'message': 'ciorna (draft) este obligatorie'}), 400

    client = _get_groq_client()
    if not client:
        return jsonify({'success': False, 'message': 'AI neconfigurat'}), 503

    try:
        system_prompt = (
            "Ești un asistent pentru recenzii de carte. "
            "Rescrie ciorna de recenzie a elevului ca o recenzie coerentă, bine scrisă, de 2-4 propoziții. "
            "Păstrează exact opinia și tonul original al elevului — nu schimba dacă e pozitivă sau negativă. "
            "Nu adăuga informații inventate. Răspunde DOAR cu textul recenziei, fără explicații."
        )
        user_prompt = f'Carte: "{titlu_carte}"\nGândul brut al elevului:\n"{ciorna}"'
        raspuns_ai = _groq_chat(client, user_prompt, system_message=system_prompt)
        return jsonify({'success': True, 'review': raspuns_ai}), 200
    except Exception:
        logger.exception('Eroare AI review-assist')
        return jsonify({'success': False, 'message': 'Cererea AI a eșuat'}), 500


@main_bp.route('/ai/book-summary/<int:carte_id>', methods=['GET'])
@jwt_required
def ai_book_summary(carte_id):
    client = _get_groq_client()
    if not client:
        return jsonify({'success': False, 'message': 'AI neconfigurat'}), 503

    try:
        book_row = db.session.execute(
            text("SELECT titlu, autor FROM carti WHERE carte_id = :id"), {'id': carte_id}
        ).fetchone()
        if not book_row:
            return jsonify({'success': False, 'message': 'Cartea nu a fost găsită'}), 404

        reviews = db.session.execute(
            text("SELECT nota, comentariu FROM recenzii WHERE carte_id = :id LIMIT 20"), {'id': carte_id}
        ).fetchall()

        if not reviews:
            return jsonify({'success': True, 'summary': '', 'no_reviews': True}), 200

        review_text = '\n'.join(f'- Nota {r[0]}/5: {r[1]}' for r in reviews)
        system_prompt = (
            "Ești un asistent de bibliotecă. Rezumă opinia generală a cititorilor despre o carte pe baza recenziilor primite. "
            "Scrie un rezumat de 2-3 propoziții. "
            "Menționează dacă e apreciată sau nu și de ce. Răspunde în română."
        )
        user_prompt = (
            f'Carte: "{book_row[0]}" de {book_row[1]}\n'
            f"Recenzii de la cititori:\n{review_text}"
        )
        raspuns_ai = _groq_chat(client, user_prompt, system_message=system_prompt)
        return jsonify({'success': True, 'summary': raspuns_ai}), 200
    except Exception:
        logger.exception('Eroare AI book-summary pentru carte_id=%d', carte_id)
        return jsonify({'success': False, 'message': 'Cererea AI a eșuat'}), 500


@main_bp.route('/ai/chat', methods=['POST'])
@jwt_required
def ai_chat():
    data = request.get_json(silent=True) or {}
    message = (data.get('message') or '').strip()
    if not message:
        return jsonify({'success': False, 'message': 'mesajul este obligatoriu'}), 400

    client = _get_groq_client()
    if not client:
        return jsonify({'success': False, 'message': 'AI neconfigurat'}), 503

    try:
        # Catalogul complet din DB
        books = db.session.execute(
            text("SELECT carte_id, titlu, autor, gen, stoc_disponibil FROM carti ORDER BY titlu")
        ).fetchall()
        book_list = '\n'.join(
            f'- [ID:{r[0]}] {r[1]} de {r[2]} (gen: {r[3]}) — {"disponibil" if r[4] > 0 else "indisponibil"}'
            for r in books
        )

        # Dacă utilizatorul e autentificat, adăugăm contextul său de citire
        user = get_current_user()
        context_utilizator = ''
        if user:
            read_rows = db.session.execute(
                text("SELECT c.titlu, c.autor FROM carti_citite cc JOIN carti c ON cc.carte_id = c.carte_id WHERE cc.user_id = :uid"),
                {'uid': user['user_id']}
            ).fetchall()
            if read_rows:
                context_utilizator = '\nCărți citite de utilizatorul curent: ' + ', '.join(f'"{r[0]}"' for r in read_rows) + '\n'

        system_prompt = (
            "Ești asistentul virtual al bibliotecii școlii CNI Suceava. "
            "Răspunzi scurt, prietenos și în română. "
            "Folosește DOAR informațiile din catalogul de mai jos — nu inventa cărți sau autori.\n\n"
            f"CATALOG COMPLET:\n{book_list}\n"
            f"{context_utilizator}"
        )
        raspuns_ai = _groq_chat(client, message, system_message=system_prompt)
        return jsonify({'success': True, 'reply': raspuns_ai}), 200
    except Exception:
        logger.exception('Eroare AI chat')
        return jsonify({'success': False, 'message': 'Cererea AI a eșuat'}), 500


# club

MAX_INVITE_HOURS = 168  # 7 zile

@main_bp.route('/club/invite', methods=['POST'])
@bibliotecar_required
def create_club_invite():
    data = request.get_json(silent=True) or {}
    try:
        expires_in_hours = int(data.get('expires_in_hours', 24))
    except (ValueError, TypeError):
        expires_in_hours = 24

    expires_in_hours = max(1, min(expires_in_hours, MAX_INVITE_HOURS))

    token = secrets.token_urlsafe(32)
    expires_at = datetime.datetime.utcnow() + datetime.timedelta(hours=expires_in_hours)
    user = request.current_user

    try:
        db.session.execute(
            text("INSERT INTO club_invites (token, created_by, expires_at) VALUES (:token, :uid, :exp)"),
            {'token': token, 'uid': user['user_id'], 'exp': expires_at}
        )
        db.session.commit()
    except Exception:
        db.session.rollback()
        logger.exception('Eroare la crearea invitației de club')
        return jsonify({'success': False, 'message': 'Eroare la generarea link-ului'}), 500

    return jsonify({
        'success': True,
        'token': token,
        'expires_at': expires_at.strftime('%d.%m.%Y %H:%M'),
        'expires_in_hours': expires_in_hours
    }), 201


@main_bp.route('/club/join/<token>', methods=['POST'])
@jwt_required
def join_club(token):
    user = request.current_user

    try:
        row = db.session.execute(
            text("SELECT expires_at FROM club_invites WHERE token = :token"),
            {'token': token}
        ).fetchone()

        if not row:
            return jsonify({'success': False, 'message': 'Link de invitație invalid'}), 404

        if datetime.datetime.utcnow() > row[0]:
            return jsonify({'success': False, 'message': 'Link-ul de invitație a expirat'}), 410

        # Verificăm dacă e deja în club
        member_row = db.session.execute(
            text("SELECT club FROM users WHERE user_id = :uid"),
            {'uid': user['user_id']}
        ).fetchone()

        if member_row and member_row[0]:
            return jsonify({'success': True, 'already_member': True, 'message': 'Ești deja membru al clubului'}), 200

        db.session.execute(
            text("UPDATE users SET club = 1 WHERE user_id = :uid"),
            {'uid': user['user_id']}
        )
        db.session.commit()
        return jsonify({'success': True, 'already_member': False, 'message': 'Bun venit în Clubul de Literatură!'}), 200
    except Exception:
        db.session.rollback()
        logger.exception('Eroare la join club pentru token %s', token)
        return jsonify({'success': False, 'message': 'Eroare la procesarea invitației'}), 500


# activitati club

def club_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        user = get_current_user()
        if not user:
            return jsonify({'success': False, 'message': 'Autentificare necesară'}), 401
        request.current_user = user
        if not user.get('club') and user.get('rol') != 'bibliotecar':
            return jsonify({'success': False, 'message': 'Acces restricționat clubului'}), 403
        return f(*args, **kwargs)
    return decorated


SAPTAMANI_VALIDE = {'anterioara', 'curenta', 'urmatoare'}
TIPURI_VALIDE    = {'anunt', 'sarcina', 'activitate'}


@main_bp.route('/club/activitati', methods=['GET'])
@club_required
def get_activitati():
    saptamana = request.args.get('saptamana', 'curenta')
    if saptamana not in SAPTAMANI_VALIDE:
        saptamana = 'curenta'
    try:
        rows = db.session.execute(
            text("""
                SELECT a.activitate_id, a.titlu, a.continut, a.tip, a.saptamana,
                       a.creat_la, a.actualizat_la, a.nume_imagine,
                       u.username, u.rol,
                       (SELECT COUNT(*) FROM comentarii_activitati c WHERE c.activitate_id = a.activitate_id) AS nr_comentarii
                FROM activitati_club a
                JOIN users u ON u.user_id = a.creat_de
                WHERE a.saptamana = :s
                ORDER BY a.creat_la DESC
            """),
            {'s': saptamana}
        ).fetchall()
        activitati = []
        for r in rows:
            activitati.append({
                'activitate_id': r[0],
                'titlu': r[1],
                'continut': r[2],
                'tip': r[3],
                'saptamana': r[4],
                'creat_la': r[5].strftime('%d.%m.%Y %H:%M') if r[5] else None,
                'actualizat_la': r[6].strftime('%d.%m.%Y %H:%M') if r[6] else None,
                'imagine_url': f'/api/club/activitati/{r[0]}/imagine' if r[7] else None,
                'autor': r[8],
                'autor_rol': r[9],
                'nr_comentarii': r[10]
            })
        return jsonify({'success': True, 'activitati': activitati}), 200
    except Exception:
        logger.exception('Eroare la get activitati club')
        return jsonify({'success': False, 'message': 'Eroare la încărcarea activităților'}), 500


@main_bp.route('/club/activitati', methods=['POST'])
@bibliotecar_required
def create_activitate():
    data = _get_request_data()
    titlu    = (data.get('titlu') or '').strip()
    continut = (data.get('continut') or '').strip()
    tip      = data.get('tip', 'activitate')
    saptamana = data.get('saptamana', 'curenta')
    image_file = request.files.get('image')

    if not titlu:
        return jsonify({'success': False, 'message': 'Titlul este obligatoriu'}), 400
    if tip not in TIPURI_VALIDE:
        tip = 'activitate'
    if saptamana not in SAPTAMANI_VALIDE:
        saptamana = 'curenta'
    if image_file and not allowed_file(image_file.filename):
        return jsonify({'success': False, 'message': 'Format imagine nepermis. Folosește png, jpg, jpeg, gif sau webp.'}), 400

    try:
        result = db.session.execute(
            text("""INSERT INTO activitati_club (titlu, continut, tip, saptamana, creat_de)
                    VALUES (:titlu, :continut, :tip, :s, :uid)"""),
            {'titlu': titlu, 'continut': continut or None, 'tip': tip,
             's': saptamana, 'uid': request.current_user['user_id']}
        )
        activitate_id = result.lastrowid
        if image_file and tip == 'anunt':
            filename = _save_club_activity_image(image_file, activitate_id)
            if filename:
                db.session.execute(
                    text("UPDATE activitati_club SET nume_imagine = :img WHERE activitate_id = :id"),
                    {'img': filename, 'id': activitate_id}
                )
        db.session.commit()
        return jsonify({'success': True, 'activitate_id': activitate_id}), 201
    except Exception:
        db.session.rollback()
        logger.exception('Eroare la creare activitate club')
        return jsonify({'success': False, 'message': 'Eroare la salvarea activității'}), 500


@main_bp.route('/club/activitati/<int:activitate_id>/imagine', methods=['GET'])
def get_activitate_image(activitate_id):
    matches = _find_files_by_prefix(CLUB_ANNOUNCEMENT_IMAGES_DIR, str(activitate_id))
    if not matches:
        return jsonify({'success': False, 'message': 'Imaginea nu a fost găsită'}), 404
    return send_from_directory(CLUB_ANNOUNCEMENT_IMAGES_DIR, matches[0])


@main_bp.route('/club/activitati/<int:activitate_id>', methods=['DELETE'])
@bibliotecar_required
def delete_activitate(activitate_id):
    try:
        for old in _find_files_by_prefix(CLUB_ANNOUNCEMENT_IMAGES_DIR, str(activitate_id)):
            try:
                os.remove(os.path.join(CLUB_ANNOUNCEMENT_IMAGES_DIR, old))
            except OSError:
                pass
        db.session.execute(
            text("DELETE FROM activitati_club WHERE activitate_id = :id"),
            {'id': activitate_id}
        )
        db.session.commit()
        return jsonify({'success': True}), 200
    except Exception:
        db.session.rollback()
        logger.exception('Eroare la ștergere activitate %s', activitate_id)
        return jsonify({'success': False, 'message': 'Eroare la ștergere'}), 500


@main_bp.route('/club/activitati/<int:activitate_id>/comentarii', methods=['GET'])
@club_required
def get_comentarii(activitate_id):
    try:
        rows = db.session.execute(
            text("""
                SELECT c.comentariu_id, c.continut, c.creat_la,
                       u.username, u.rol, u.user_id
                FROM comentarii_activitati c
                JOIN users u ON u.user_id = c.user_id
                WHERE c.activitate_id = :id
                ORDER BY c.creat_la ASC
            """),
            {'id': activitate_id}
        ).fetchall()
        comentarii = [{
            'comentariu_id': r[0],
            'continut': r[1],
            'creat_la': r[2].strftime('%d.%m.%Y %H:%M') if r[2] else None,
            'autor': r[3],
            'autor_rol': r[4],
            'user_id': r[5]
        } for r in rows]
        return jsonify({'success': True, 'comentarii': comentarii}), 200
    except Exception:
        logger.exception('Eroare la get comentarii activitate %s', activitate_id)
        return jsonify({'success': False, 'message': 'Eroare la încărcarea comentariilor'}), 500


@main_bp.route('/club/activitati/<int:activitate_id>/comentarii', methods=['POST'])
@club_required
def post_comentariu(activitate_id):
    data = request.get_json(silent=True) or {}
    continut = (data.get('continut') or '').strip()
    if not continut:
        return jsonify({'success': False, 'message': 'Comentariul nu poate fi gol'}), 400
    if len(continut) > 2000:
        return jsonify({'success': False, 'message': 'Comentariul este prea lung (max 2000 caractere)'}), 400

    # Verificăm că activitatea există
    exists = db.session.execute(
        text("SELECT 1 FROM activitati_club WHERE activitate_id = :id"),
        {'id': activitate_id}
    ).fetchone()
    if not exists:
        return jsonify({'success': False, 'message': 'Activitatea nu există'}), 404

    try:
        result = db.session.execute(
            text("INSERT INTO comentarii_activitati (activitate_id, user_id, continut) VALUES (:aid, :uid, :c)"),
            {'aid': activitate_id, 'uid': request.current_user['user_id'], 'c': continut}
        )
        db.session.commit()
        return jsonify({'success': True, 'comentariu_id': result.lastrowid}), 201
    except Exception:
        db.session.rollback()
        logger.exception('Eroare la post comentariu activitate %s', activitate_id)
        return jsonify({'success': False, 'message': 'Eroare la salvarea comentariului'}), 500


@main_bp.route('/club/activitati/<int:activitate_id>/comentarii/<int:comentariu_id>', methods=['DELETE'])
@club_required
def delete_comentariu(activitate_id, comentariu_id):
    user = request.current_user
    row = db.session.execute(
        text("SELECT user_id FROM comentarii_activitati WHERE comentariu_id = :cid AND activitate_id = :aid"),
        {'cid': comentariu_id, 'aid': activitate_id}
    ).fetchone()
    if not row:
        return jsonify({'success': False, 'message': 'Comentariul nu există'}), 404
    if row[0] != user['user_id'] and user.get('rol') != 'bibliotecar':
        return jsonify({'success': False, 'message': 'Nu poți șterge comentariul altcuiva'}), 403
    try:
        db.session.execute(
            text("DELETE FROM comentarii_activitati WHERE comentariu_id = :cid"),
            {'cid': comentariu_id}
        )
        db.session.commit()
        return jsonify({'success': True}), 200
    except Exception:
        db.session.rollback()
        logger.exception('Eroare la ștergere comentariu %s', comentariu_id)
        return jsonify({'success': False, 'message': 'Eroare la ștergere'}), 500

# --- CLUB THREADS ---

@main_bp.route('/club/threads', methods=['GET'])
@club_required
def get_club_threads():
    try:
        threads = db.session.query(
            ClubThreads.thread_id,
            ClubThreads.titlu,
            ClubThreads.continut,
            ClubThreads.creat_la,
            Users.username,
            Users.user_id.label('autor_id')
        ).join(Users, ClubThreads.creat_de == Users.user_id).order_by(ClubThreads.creat_la.desc()).all()
        
        result = []
        for t in threads:
            result.append({
                'thread_id': t.thread_id,
                'titlu': t.titlu,
                'continut': t.continut,
                'creat_la': t.creat_la.isoformat() if t.creat_la else None,
                'autor': t.username,
                'autor_id': t.autor_id
            })
        return jsonify({'threads': result}), 200
    except Exception as e:
        logger.error(f"Eroare la preluarea thread-urilor: {e}")
        return jsonify({'error': 'A apărut o eroare la preluarea thread-urilor.'}), 500

@main_bp.route('/club/threads', methods=['POST'])
@club_required
def create_club_thread():
    data = request.get_json()
    if not data or not data.get('titlu') or not data.get('continut'):
        return jsonify({'error': 'Titlul și conținutul sunt obligatorii.'}), 400
        
    try:
        new_thread = ClubThreads(
            titlu=data['titlu'],
            continut=data['continut'],
            creat_de=request.current_user.user_id
        )
        db.session.add(new_thread)
        db.session.commit()
        return jsonify({'message': 'Thread creat cu succes.', 'thread_id': new_thread.thread_id}), 201
    except Exception as e:
        db.session.rollback()
        logger.error(f"Eroare la crearea thread-ului: {e}")
        return jsonify({'error': 'A apărut o eroare la crearea thread-ului.'}), 500

@main_bp.route('/club/threads/<int:thread_id>', methods=['GET'])
@club_required
def get_club_thread(thread_id):
    try:
        thread = db.session.query(
            ClubThreads.thread_id,
            ClubThreads.titlu,
            ClubThreads.continut,
            ClubThreads.creat_la,
            Users.username,
            Users.user_id.label('autor_id')
        ).join(Users, ClubThreads.creat_de == Users.user_id).filter(ClubThreads.thread_id == thread_id).first()
        
        if not thread:
            return jsonify({'error': 'Thread-ul nu a fost găsit.'}), 404
            
        comments_query = db.session.query(
            ThreadComments.comentariu_id,
            ThreadComments.continut,
            ThreadComments.likes,
            ThreadComments.creat_la,
            Users.username,
            Users.user_id.label('autor_id')
        ).join(Users, ThreadComments.user_id == Users.user_id).filter(ThreadComments.thread_id == thread_id).order_by(ThreadComments.creat_la.asc()).all()
        
        comments = []
        for c in comments_query:
            subcomments_query = db.session.query(
                ThreadSubcomments.subcomentariu_id,
                ThreadSubcomments.continut,
                ThreadSubcomments.likes,
                ThreadSubcomments.creat_la,
                Users.username,
                Users.user_id.label('autor_id')
            ).join(Users, ThreadSubcomments.user_id == Users.user_id).filter(ThreadSubcomments.comentariu_id == c.comentariu_id).order_by(ThreadSubcomments.creat_la.asc()).all()
            
            subcomments = []
            for sc in subcomments_query:
                subcomments.append({
                    'subcomentariu_id': sc.subcomentariu_id,
                    'continut': sc.continut,
                    'likes': sc.likes,
                    'creat_la': sc.creat_la.isoformat() if sc.creat_la else None,
                    'autor': sc.username,
                    'autor_id': sc.autor_id
                })
                
            comments.append({
                'comentariu_id': c.comentariu_id,
                'continut': c.continut,
                'likes': c.likes,
                'creat_la': c.creat_la.isoformat() if c.creat_la else None,
                'autor': c.username,
                'autor_id': c.autor_id,
                'subcomments': subcomments
            })
            
        return jsonify({
            'thread': {
                'thread_id': thread.thread_id,
                'titlu': thread.titlu,
                'continut': thread.continut,
                'creat_la': thread.creat_la.isoformat() if thread.creat_la else None,
                'autor': thread.username,
                'autor_id': thread.autor_id,
                'comments': comments
            }
        }), 200
        
    except Exception as e:
        logger.error(f"Eroare la preluarea thread-ului: {e}")
        return jsonify({'error': 'A apărut o eroare la preluarea thread-ului.'}), 500

@main_bp.route('/club/threads/<int:thread_id>/comments', methods=['POST'])
@club_required
def add_thread_comment(thread_id):
    data = request.get_json()
    if not data or not data.get('continut'):
        return jsonify({'error': 'Conținutul este obligatoriu.'}), 400
        
    try:
        new_comment = ThreadComments(
            thread_id=thread_id,
            user_id=request.current_user.user_id,
            continut=data['continut']
        )
        db.session.add(new_comment)
        db.session.commit()
        return jsonify({'message': 'Comentariu adăugat cu succes.', 'comentariu_id': new_comment.comentariu_id}), 201
    except Exception as e:
        db.session.rollback()
        logger.error(f"Eroare la adăugarea comentariului: {e}")
        return jsonify({'error': 'A apărut o eroare la adăugarea comentariului.'}), 500

@main_bp.route('/club/threads/comments/<int:comentariu_id>/subcomments', methods=['POST'])
@club_required
def add_thread_subcomment(comentariu_id):
    data = request.get_json()
    if not data or not data.get('continut'):
        return jsonify({'error': 'Conținutul este obligatoriu.'}), 400
        
    try:
        new_subcomment = ThreadSubcomments(
            comentariu_id=comentariu_id,
            user_id=request.current_user.user_id,
            continut=data['continut']
        )
        db.session.add(new_subcomment)
        db.session.commit()
        return jsonify({'message': 'Subcomentariu adăugat cu succes.', 'subcomentariu_id': new_subcomment.subcomentariu_id}), 201
    except Exception as e:
        db.session.rollback()
        logger.error(f"Eroare la adăugarea subcomentariului: {e}")
        return jsonify({'error': 'A apărut o eroare la adăugarea subcomentariului.'}), 500

@main_bp.route('/club/threads/comments/<int:comentariu_id>/like', methods=['POST'])
@club_required
def like_thread_comment(comentariu_id):
    try:
        comment = db.session.query(ThreadComments).filter_by(comentariu_id=comentariu_id).first()
        if not comment:
            return jsonify({'error': 'Comentariul nu a fost găsit.'}), 404
            
        comment.likes += 1
        db.session.commit()
        return jsonify({'message': 'Comentariu apreciat cu succes.', 'likes': comment.likes}), 200
    except Exception as e:
        db.session.rollback()
        logger.error(f"Eroare la aprecierea comentariului: {e}")
        return jsonify({'error': 'A apărut o eroare la aprecierea comentariului.'}), 500

@main_bp.route('/club/threads/subcomments/<int:subcomentariu_id>/like', methods=['POST'])
@club_required
def like_thread_subcomment(subcomentariu_id):
    try:
        subcomment = db.session.query(ThreadSubcomments).filter_by(subcomentariu_id=subcomentariu_id).first()
        if not subcomment:
            return jsonify({'error': 'Subcomentariul nu a fost găsit.'}), 404
            
        subcomment.likes += 1
        db.session.commit()
        return jsonify({'message': 'Subcomentariu apreciat cu succes.', 'likes': subcomment.likes}), 200
    except Exception as e:
        db.session.rollback()
        logger.error(f"Eroare la aprecierea subcomentariului: {e}")
        return jsonify({'error': 'A apărut o eroare la aprecierea subcomentariului.'}), 500
