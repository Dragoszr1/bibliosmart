from flask import request, jsonify, send_file
from sqlalchemy import text
from app.database import db
from app.utils.validators import normalize_role
import datetime
import io
import logging
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

def librarian_report_docx():
    try:
        students = db.session.execute(
            text("SELECT user_id, username, email FROM users WHERE rol NOT IN ('1','bibliotecar','administrator') ORDER BY username ASC")
        ).fetchall()

        read_rows = db.session.execute(
            text(
                "SELECT cc.user_id, c.titlu, c.autor "
                "FROM carti_citite cc JOIN carti c ON cc.carte_id = c.carte_id"
            )
        ).fetchall()
        read_map = {}
        for r in read_rows:
            read_map.setdefault(r[0], []).append((r[1], r[2]))

        borrow_rows = db.session.execute(
            text(
                "SELECT ia.user_id, c.titlu, c.autor, ia.data_imprumut, ia.data_scadenta "
                "FROM imprumuturi_active ia JOIN carti c ON ia.carte_id = c.carte_id "
                "ORDER BY ia.data_scadenta ASC"
            )
        ).fetchall()
        borrow_map = {}
        for r in borrow_rows:
            borrow_map.setdefault(r[0], []).append((r[1], r[2], r[3], r[4]))

        approved_rows = db.session.execute(
            text(
                "SELECT cr.user_id, c.titlu, c.autor, cr.ridicare_de_la, cr.ridicare_pana_la "
                "FROM cereri_carti cr JOIN carti c ON cr.carte_id = c.carte_id "
                "WHERE cr.status = 'approved' "
                "ORDER BY cr.updated_at ASC"
            )
        ).fetchall()
        approved_map = {}
        for r in approved_rows:
            approved_map.setdefault(r[0], []).append((r[1], r[2], r[3], r[4]))

        doc = Document()

        title_para = doc.add_heading('Raport Bibliotecă – Elevi și Împrumuturi', level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        generated_para = doc.add_paragraph(
            f'Generat la: {datetime.datetime.now().strftime("%d.%m.%Y %H:%M")}'
        )
        generated_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        if not students:
            doc.add_paragraph('Nu există elevi înregistrați.')
        else:
            for uid, username, email in students:
                doc.add_heading(username, level=1)
                email_para = doc.add_paragraph()
                email_run = email_para.add_run(f'Email: {email}')
                email_run.italic = True

                approved = approved_map.get(uid, [])
                doc.add_heading('Aprobate — în așteptarea ridicării', level=2)
                if approved:
                    tbl0 = doc.add_table(rows=1, cols=4)
                    tbl0.style = 'Table Grid'
                    hdr0 = tbl0.rows[0].cells
                    hdr0[0].text = 'Titlu'; hdr0[1].text = 'Autor'
                    hdr0[2].text = 'Interval ridicare (de la)'; hdr0[3].text = 'Interval ridicare (până la)'
                    for cell in hdr0:
                        cell.paragraphs[0].runs[0].bold = True
                    for titlu, autor, pf, pu in approved:
                        rc = tbl0.add_row().cells
                        rc[0].text = titlu; rc[1].text = autor
                        rc[2].text = pf.strftime('%d.%m.%Y %H:%M') if pf else '—'
                        rc[3].text = pu.strftime('%d.%m.%Y %H:%M') if pu else '—'
                else:
                    doc.add_paragraph('Nicio carte în așteptarea ridicării.')

                borrows = borrow_map.get(uid, [])
                doc.add_heading('Cărți împrumutate activ', level=2)
                if borrows:
                    tbl = doc.add_table(rows=1, cols=4)
                    tbl.style = 'Table Grid'
                    hdr = tbl.rows[0].cells
                    hdr[0].text = 'Titlu'
                    hdr[1].text = 'Autor'
                    hdr[2].text = 'Data împrumutului'
                    hdr[3].text = 'Dată scadentă (30 zile)'
                    for cell in hdr:
                        run = cell.paragraphs[0].runs[0]
                        run.bold = True
                    for titlu, autor, data_imp, data_sc in borrows:
                        row_cells = tbl.add_row().cells
                        row_cells[0].text = titlu
                        row_cells[1].text = autor
                        row_cells[2].text = data_imp.strftime('%d.%m.%Y') if data_imp else '—'
                        row_cells[3].text = data_sc.strftime('%d.%m.%Y') if data_sc else '—'
                else:
                    doc.add_paragraph('Niciun împrumut activ.')

                reads = read_map.get(uid, [])
                doc.add_heading('Cărți citite', level=2)
                if reads:
                    tbl2 = doc.add_table(rows=1, cols=2)
                    tbl2.style = 'Table Grid'
                    hdr2 = tbl2.rows[0].cells
                    hdr2[0].text = 'Titlu'
                    hdr2[1].text = 'Autor'
                    for cell in hdr2:
                        run = cell.paragraphs[0].runs[0]
                        run.bold = True
                    for titlu, autor in reads:
                        row_cells = tbl2.add_row().cells
                        row_cells[0].text = titlu
                        row_cells[1].text = autor
                else:
                    doc.add_paragraph('Nicio carte citită înregistrată.')

                doc.add_paragraph()

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        filename = f'raport_biblioteca_{datetime.datetime.now().strftime("%Y%m%d_%H%M")}.docx'
        return send_file(
            buf,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception:
        logger.exception('Eroare la generarea raportului Word')
        return jsonify({'success': False, 'message': 'Eroare la generarea raportului'}), 500

def admin_get_users():
    try:
        rows = db.session.execute(
            text("SELECT user_id, username, email, rol, telefon FROM users ORDER BY username ASC")
        ).fetchall()
        users = [
            {
                'user_id': r[0],
                'username': r[1],
                'email': r[2],
                'rol': normalize_role(r[3]),
                'telefon': r[4]
            }
            for r in rows
        ]
        return jsonify({'success': True, 'users': users}), 200
    except Exception:
        logger.exception('Eroare la preluarea listei de utilizatori')
        return jsonify({'success': False, 'error': 'Eroare la preluarea utilizatorilor'}), 500

def admin_get_user_detail(user_id):
    try:
        user_row = db.session.execute(
            text("SELECT user_id, username, email, rol, telefon, description FROM users WHERE user_id = :uid"),
            {'uid': user_id}
        ).fetchone()
        if not user_row:
            return jsonify({'success': False, 'message': 'Utilizatorul nu a fost găsit'}), 404

        borrowed_rows = db.session.execute(
            text("""
                SELECT ia.imprumut_id, c.carte_id, c.titlu, c.autor, c.ISBN,
                       ia.data_imprumut, ia.data_scadenta
                FROM imprumuturi_active ia
                JOIN carti c ON ia.carte_id = c.carte_id
                WHERE ia.user_id = :uid
                ORDER BY ia.data_imprumut DESC
            """),
            {'uid': user_id}
        ).fetchall()

        read_rows = db.session.execute(
            text("""
                SELECT c.carte_id, c.titlu, c.autor, c.ISBN
                FROM carti_citite cc
                JOIN carti c ON cc.carte_id = c.carte_id
                WHERE cc.user_id = :uid
            """),
            {'uid': user_id}
        ).fetchall()

        history_rows = db.session.execute(
            text("""
                SELECT cc.cerere_id, c.titlu, c.autor, cc.status, cc.created_at,
                       cc.ridicare_de_la, cc.ridicare_pana_la
                FROM cereri_carti cc
                JOIN carti c ON cc.carte_id = c.carte_id
                WHERE cc.user_id = :uid
                ORDER BY cc.created_at DESC
            """),
            {'uid': user_id}
        ).fetchall()

        review_rows = db.session.execute(
            text("""
                SELECT r.id, c.titlu, c.autor, r.nota, r.comentariu
                FROM recenzii r
                JOIN carti c ON r.carte_id = c.carte_id
                WHERE r.user_id = :uid
                ORDER BY r.id DESC
            """),
            {'uid': user_id}
        ).fetchall()

        return jsonify({
            'success': True,
            'user': {
                'user_id': user_row[0],
                'username': user_row[1],
                'email': user_row[2],
                'rol': normalize_role(user_row[3]),
                'telefon': user_row[4],
                'description': user_row[5]
            },
            'books_borrowed': [
                {
                    'imprumut_id': r[0],
                    'carte_id': r[1], 'titlu': r[2], 'autor': r[3], 'ISBN': r[4],
                    'borrowed_at': r[5].isoformat() if r[5] else None,
                    'due_at': r[6].isoformat() if r[6] else None
                }
                for r in borrowed_rows
            ],
            'books_read': [
                {'carte_id': r[0], 'titlu': r[1], 'autor': r[2], 'ISBN': r[3]}
                for r in read_rows
            ],
            'borrow_history': [
                {
                    'cerere_id': r[0], 'titlu': r[1], 'autor': r[2],
                    'status': r[3],
                    'created_at': r[4].isoformat() if r[4] else None,
                    'ridicare_de_la': r[5].strftime('%d.%m.%Y %H:%M') if r[5] else None,
                    'ridicare_pana_la': r[6].strftime('%d.%m.%Y %H:%M') if r[6] else None
                }
                for r in history_rows
            ],
            'reviews': [
                {'id': r[0], 'titlu': r[1], 'autor': r[2], 'nota': r[3], 'comentariu': r[4]}
                for r in review_rows
            ]
        }), 200
    except Exception:
        logger.exception('Eroare la preluarea detaliilor utilizatorului %d', user_id)
        return jsonify({'success': False, 'error': 'Eroare la preluarea detaliilor utilizatorului'}), 500

def prelungeste_imprumut(imprumut_id):
    data = request.get_json(silent=True) or {}
    data_noua = (data.get('data_scadenta') or '').strip()
    if not data_noua:
        return jsonify({'success': False, 'message': 'data_scadenta este obligatorie'}), 400
    try:
        data_scadenta = datetime.datetime.fromisoformat(data_noua)
    except ValueError:
        return jsonify({'success': False, 'message': 'Format dată invalid. Folosește YYYY-MM-DDTHH:MM'}), 400
    try:
        result = db.session.execute(
            text("UPDATE imprumuturi_active SET data_scadenta = :ds WHERE imprumut_id = :id"),
            {'ds': data_scadenta, 'id': imprumut_id}
        )
        if result.rowcount == 0:
            return jsonify({'success': False, 'message': 'Împrumutul nu a fost găsit'}), 404
        db.session.commit()
        return jsonify({'success': True, 'message': 'Data scadentă actualizată'}), 200
    except Exception:
        db.session.rollback()
        logger.exception('Eroare la prelungire împrumut %d', imprumut_id)
        return jsonify({'success': False, 'message': 'Eroare la actualizare'}), 500

def returneaza_imprumut(imprumut_id):
    try:
        row = db.session.execute(
            text("SELECT user_id, carte_id FROM imprumuturi_active WHERE imprumut_id = :id"),
            {'id': imprumut_id}
        ).fetchone()
        if not row:
            return jsonify({'success': False, 'message': 'Împrumutul nu a fost găsit'}), 404
        user_id, carte_id = row

        db.session.execute(
            text("DELETE FROM imprumuturi_active WHERE imprumut_id = :id"),
            {'id': imprumut_id}
        )
        db.session.execute(
            text("UPDATE carti SET stoc_disponibil = stoc_disponibil + 1, "
                 "imprumutat = (stoc_disponibil + 1 < stoc_total) "
                 "WHERE carte_id = :cid"),
            {'cid': carte_id}
        )

        already = db.session.execute(
            text("SELECT 1 FROM carti_citite WHERE user_id = :uid AND carte_id = :cid"),
            {'uid': user_id, 'cid': carte_id}
        ).fetchone()
        if not already:
            db.session.execute(
                text("INSERT INTO carti_citite (user_id, carte_id) VALUES (:uid, :cid)"),
                {'uid': user_id, 'cid': carte_id}
            )

        db.session.commit()
        return jsonify({'success': True, 'message': 'Carte returnată cu succes'}), 200
    except Exception:
        db.session.rollback()
        logger.exception('Eroare la returnarea împrumutului %d', imprumut_id)
        return jsonify({'success': False, 'message': 'Eroare la returnare'}), 500
