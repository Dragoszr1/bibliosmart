import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()
doc.add_heading('Documentație Tehnică Completă - Platforma Bibliosmart', 0)

# Intro
doc.add_paragraph('Această documentație acoperă în detaliu toate nivelurile aplicației (frontend, backend, bază de date, server web, optimizări, caching, scalabilitate și securitate), exemplificând direct cu secvențe din codul sursă real al platformei.')

# 1. Metodologie și Arhitectură
doc.add_heading('1. Metodologie și Arhitectură Generală', level=1)
doc.add_paragraph('Aplicația urmează o arhitectură "Client-Server" puternic decuplată. Frontend-ul este un SPA (Single Page Application) construit în Vue 3 (Composition API) și Vite. Asta înseamnă că browserul descarcă un pachet inițial extrem de mic de fișiere statice (HTML/JS/CSS), iar navigarea prin platformă se face prin "vue-router" fără a face vreodată reîncărcarea (refresh) ferestrei de browser.')
doc.add_paragraph('Backend-ul este dezvoltat în Python cu microframework-ul Flask, funcționând exclusiv ca un API RESTful care servește JSON-uri. Toate interogările se fac prin `fetch()`.')
doc.add_paragraph('Această decuplare aduce mari beneficii pe partea de scalabilitate: componentele sunt independente, iar backend-ul (fiind stateless prin JWT) poate fi replicat în spatele unui Load Balancer foarte ușor.')

# 2. Securitate: Autentificare JWT și HttpOnly Cookies
doc.add_heading('2. Securitate: JWT și Cookie-uri HttpOnly', level=1)
doc.add_paragraph('Spre deosebire de o arhitectură clasică nesigură unde token-ul de login s-ar salva în LocalStorage, platforma Bibliosmart folosește cookie-uri HttpOnly pentru protecție totală împotriva atacurilor XSS (Cross-Site Scripting). Javascript-ul din browserul potențialului atacator nu poate accesa acest cookie.')
doc.add_paragraph('Secvența de mai jos (preluată direct din backend/app/utils/auth.py) arată cum se setează acest cookie în mod securizat:')

code1 = """def set_jwt_cookie(response, token):
    response.set_cookie(
        current_app.config['JWT_COOKIE_NAME'],
        token,
        httponly=True,  # Restricționează complet accesul la cookie din JavaScript
        secure=current_app.config['JWT_COOKIE_SECURE'], # Flag activ în producție pentru HTTPS exclusiv
        samesite=current_app.config['JWT_COOKIE_SAMESITE'], # Setat pe 'Lax' pentru protecție CSRF
        max_age=int(current_app.config['JWT_ACCESS_TOKEN_EXPIRES'].total_seconds()),
        path='/'
    )
    return response"""
p = doc.add_paragraph()
r = p.add_run(code1)
r.font.name = 'Courier New'
r.font.size = Pt(9)

doc.add_paragraph('Deoarece folosim cookie-uri securizate (care ascund tokenul JWT), frontend-ul trebuie să adauge explicit un flag la absolut orice apel API către backend. Exemplu real din frontend (frontend/src/pages/Club.vue):')
code2 = """const res = await fetch('/api/auth/me', { credentials: 'include' })"""
p = doc.add_paragraph()
r = p.add_run(code2)
r.font.name = 'Courier New'
r.font.size = Pt(10)

# 3. Securitate: SQL Injection și Validări
doc.add_heading('3. Securitate: Protecția contra SQL Injection', level=1)
doc.add_paragraph('Baza de date (MySQL) este interogată folosind biblioteca SQLAlchemy (ca Layer de execuție și ORM). Orice input trimis de utilizator din frontend este trecut obligatoriu prin interogări "Prepared/Parameterized" folosind metoda text() și un dicționar de argumente.')
doc.add_paragraph('Datorită parametrilor numiți (:username, :email), motorul bazei de date preia datele direct ca valori pur textuale (String), nelăsând absolut nicio posibilitate pentru execuția codului SQL malițios inserat intenționat. Fragment din backend/app/controllers/auth_controller.py:')
code3 = """    insert_query = text(
        "INSERT INTO users (username, email, hashed_password, rol, telefon, club) "
        "VALUES (:username, :email, :hashed_password, :rol, :telefon, :club)"
    )
    # Valorile sunt injectate sigur prin dicționar
    db.session.execute(insert_query, values)"""
p = doc.add_paragraph()
r = p.add_run(code3)
r.font.name = 'Courier New'
r.font.size = Pt(9)

# 4. Securitate: Headere HTTP Nginx & Flask
doc.add_heading('4. Securitate: Headere HTTP Defensiva', level=1)
doc.add_paragraph('Backend-ul este configurat să emită capete (headers) HTTP stricte pe absolut fiecare request primit, interceptat prin decoratorul @app.after_request (din backend/app/__init__.py):')
code4 = """    @app.after_request
    def add_security_headers(response):
        response.headers['X-Content-Type-Options'] = 'nosniff' # Oprește browserul să ghicească tipul MIME malițios
        response.headers['X-Frame-Options'] = 'DENY' # Previne Clickjacking (încorporarea site-ului în iframe-uri invizibile)
        response.headers['Referrer-Policy'] = 'strict-origin-when-cross-origin'
        response.headers['Content-Security-Policy'] = "default-src 'self' data: blob:; frame-ancestors 'none';"
        return response"""
p = doc.add_paragraph()
r = p.add_run(code4)
r.font.name = 'Courier New'
r.font.size = Pt(9)

# 5. Caching și Frontend Optimizations
doc.add_heading('5. Caching Avansat & Frontend (LocalForage și Vite)', level=1)
doc.add_paragraph('Aplicația Vue.js folosește Vite. Compilatorul Vite realizează "chunk splitting" automat și minifică (șterge spațiile și optimizează JS/CSS) codul. Acest proces garantează un TTFB (Time to First Byte) excelent și performanță maximă.')
doc.add_paragraph('La nivel logic, site-ul implementează "Caching agresiv în browser" (Stale-While-Revalidate) folosind LocalForage (o abstractizare asincronă peste baza de date nativă din browser, IndexedDB).')
doc.add_paragraph('Când un utilizator deschide lista de cărți, frontend-ul încarcă INSTANTANEU datele stocate în browser de la ultima vizită, iar într-un thread secundar interoghează serverul pe fundal pentru date fresh. Exemplu exact din frontend/src/pages/Books.vue:')
code5 = """        // 1. Încercăm să luăm datele din cache (localforage/IndexedDB) mai întâi pentru afișare instantanee
        const cachedBooks = await localforage.getItem('cachedBookList');
        if (cachedBooks) {
          this.books = cachedBooks;
          this.initialLoading = false; // Afișează cărțile cached imediat (0ms latență pe rețea!)
        }
        
        // 2. În fundal, aduce cărțile fresh
        const response = await fetch('/api/books');
        const data = await response.json();
        this.books = data.books; // Suprascrie datele vechi cu cele noi
        
        // 3. Actualizează cache-ul din browser pentru următoarea vizită
        await localforage.setItem('cachedBookList', data.books);"""
p = doc.add_paragraph()
r = p.add_run(code5)
r.font.name = 'Courier New'
r.font.size = Pt(9)
doc.add_paragraph('Această tehnică oferă utilizatorului iluzia unui timp de răspuns de rețea practic egal cu zero.')

# 6. Optimizări Backend & Baza de Date
doc.add_heading('6. Optimizări Backend (Connection Pooling)', level=1)
doc.add_paragraph('Pentru a nu supraîncărca memoria, rutele grele folosesc limitări la nivel de SQL (ex. "SELECT ... LIMIT 15" în `get_recent_books()`).')
doc.add_paragraph('Mai mult, SQLAlchemy implementează un mecanism avansat de Connection Pooling. Datorită flag-ului `pool_pre_ping: True`, sistemul testează în background conexiunea la MySQL de fiecare dată înainte de a face un query. Dacă serverul MySQL a omorât conexiunea din cauza inactivității prelungite (overnight timeout), sistemul Flask o regenerează complet automat, prevenind blocajul letal de tip "MySQL server has gone away". (sursa backend/config.py):')
code6 = """    # SQLAlchemy Pool Options to prevent "MySQL server has gone away"
    SQLALCHEMY_ENGINE_OPTIONS = {
        'pool_pre_ping': True,
        'pool_recycle': 280,
    }"""
p = doc.add_paragraph()
r = p.add_run(code6)
r.font.name = 'Courier New'
r.font.size = Pt(10)

# 7. Scalabilitate și Reverse Proxy (Nginx)
doc.add_heading('7. Arhitectura de Server (Nginx Reverse Proxy & Load Balancing)', level=1)
doc.add_paragraph('Configurația de server central se bazează pe performanța uriașă a Nginx-ului (extras din fișierul biblioteca.nginx.conf). Nginx interceptează tot traficul web de pe porturile 80/443. Pentru resursele statice Vue (din `frontend/dist`), le returnează instant din memorie. Pentru că e un SPA, activează "fallback loop" pe /index.html. Când interceptează prefixul `/api/`, el rutează mai departe în background traficul către procesul Flask/Gunicorn pe portul 8000.')
code7 = """    # SPA fallback — orice altceva duce în interfața Vue
    location / {
        try_files $uri $uri/ /index.html;
    }

    # Proxy API requests to Flask/Gunicorn
    location /api/ {
        proxy_pass         http://127.0.0.1:8000;
        proxy_set_header   Host $host;
        proxy_set_header   X-Real-IP $remote_addr;
        proxy_set_header   X-Forwarded-For $proxy_add_x_forwarded_for;
        
        # Required for file uploads (mărește limita pentru PDF-uri și imagini mari)
        client_max_body_size 10M;
    }"""
p = doc.add_paragraph()
r = p.add_run(code7)
r.font.name = 'Courier New'
r.font.size = Pt(9)
doc.add_paragraph('De ce este această arhitectură de "Scalabilitate extremă"? Deoarece aplicația Backend Flask nu reține starea niciunei logări din memorie (logările fiind Stateless prin Tokenul JWT codificat în cookie). Astfel, la o invazie de trafic uriaș, administratorul poate pur și simplu să ruleze 5-10 copii ale serverului Flask (pe porturi precum 8001, 8002...), iar în blocul `proxy_pass` din Nginx să adauge un flux de "upstream round-robin" care va distribui traficul vizitatorilor perfect egal pe acele servere. Sistemul e perfect pregătit de Load Balancing.')

doc.save('C:/Users/zimbr/biblioteca-repo/Documentatie_Completa_Cod.docx')
print('Document complet generat cu succes!')
