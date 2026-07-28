from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
doc.add_heading('Documentație Tehnică Detaliată - Platforma Bibliosmart', 0)

# Section 1
doc.add_heading('1. Metodologie și Arhitectură Generală', level=1)
p = doc.add_paragraph('Platforma Bibliosmart este dezvoltată pe baza unei arhitecturi moderne Client-Server de tip SPA (Single Page Application). Metodologia implică o decuplare completă între Frontend și Backend, comunicarea realizându-se exclusiv prin intermediul unui API RESTful, iar schimbul de date se face în format JSON.\n\n')
p.add_run('Tehnologii principale folosite:').bold = True
doc.add_paragraph('Frontend: Vue 3, Vite, Tailwind CSS, Vue Router, Axios.', style='List Bullet')
doc.add_paragraph('Backend: Python, Flask, SQLAlchemy.', style='List Bullet')
doc.add_paragraph('Bază de date: MySQL.', style='List Bullet')
doc.add_paragraph('Server Web / Reverse Proxy: Nginx.', style='List Bullet')

# Section 2
doc.add_heading('2. Backend (Flask + MySQL)', level=1)
doc.add_paragraph('Backend-ul respectă un șablon de proiectare modular (asemanator MVC). Rutele (Routes) preiau request-urile HTTP, Controllerele conțin logica de business, iar Modelele (Models) definesc schemele bazei de date prin SQLAlchemy ORM.')
doc.add_paragraph('Pentru a asigura stabilitatea conexiunii cu baza de date (evitând eroarea "MySQL server has gone away"), SQLAlchemy este configurat cu un sistem de connection pooling.')
doc.add_paragraph('Exemplu de configurare a conexiunii (config.py):', style='Normal')
code = doc.add_paragraph()
code.add_run('''    SQLALCHEMY_ENGINE_OPTIONS = {
        'pool_pre_ping': True,
        'pool_recycle': 280,
    }''').font.name = 'Courier New'

doc.add_paragraph('Sectiunea de controller asigura validarea stricta a datelor la adaugarea si manipularea informatiilor (ex: validarea lungimii, formatului ISBN), prevenind erorile si inserțiile incorecte:')
code = doc.add_paragraph()
code.add_run('''    if not titlu or len(titlu) > 50:
        return jsonify({'success': False, 'message': 'Titlul nu poate fi gol și trebuie să aibă maximum 50 de caractere'}), 400
    if not isbn or len(isbn) not in (10, 13) or not isbn.isalnum():
        return jsonify({'success': False, 'message': 'ISBN-ul trebuie să aibă 10 sau 13 caractere alfanumerice'}), 400''').font.name = 'Courier New'

# Section 3
doc.add_heading('3. Frontend (Vue 3 + Vite)', level=1)
doc.add_paragraph('Interfata cu utilizatorul este un SPA reactiv. Datorita Vue 3 si Vite, aplicatia se incarca instant, iar viteza de dezvoltare este optimizata prin Hot Module Replacement (HMR). Design-ul vizual (UI/UX) este implementat cu Tailwind CSS, asigurand un aspect modern si adaptabil (responsive/mobile-first).')
doc.add_paragraph('Starea aplicatiei si request-urile HTTP sunt gestionate de Axios. Interactiunile asincrone asigura ca pagina nu se reincarca niciodata, oferind o experienta fluida (app-like).')
doc.add_paragraph('Configurarea Axios pentru a include cookie-urile de autentificare in mod securizat:')
code = doc.add_paragraph()
code.add_run('''axios.defaults.withCredentials = true;
// Request-urile către API vor include automat JWT-ul stocat în HttpOnly cookie.''').font.name = 'Courier New'

# Section 4
doc.add_heading('4. Securitate', level=1)
doc.add_paragraph('Securitatea este un pilon central al platformei, fiind abordata pe mai multe straturi:', style='Normal')
p = doc.add_paragraph()
p.add_run('Autentificare JWT în HttpOnly Cookies:').bold = True
p.add_run(' Token-urile JWT nu sunt stocate in LocalStorage (unde ar fi vulnerabile), ci in cookie-uri setate ca HttpOnly. Acest lucru previne atacurile de tip XSS (Cross-Site Scripting). Cookie-urile au si flag-urile Secure si SameSite="Lax".')
p = doc.add_paragraph()
p.add_run('SQL Injection:').bold = True
p.add_run(' Utilizarea ORM-ului SQLAlchemy si a block-urilor de text prepared (text() paramterizat) protejeaza complet baza de date impotriva injecțiilor SQL.')
code = doc.add_paragraph()
code.add_run('''    insert_query = text(
        "INSERT INTO carti (titlu, autor, ISBN, stoc_total) "
        "VALUES (:titlu, :autor, :ISBN, :stoc_total)"
    )
    db.session.execute(insert_query, {'titlu': titlu, 'autor': autor, 'ISBN': isbn, 'stoc_total': stoc_total})''').font.name = 'Courier New'

p = doc.add_paragraph()
p.add_run('Security Headers (Flask & Nginx):').bold = True
p.add_run(' Backend-ul adauga headere de securitate la fiecare request pentru a preveni atacurile de tip clickjacking si sniffing.')
code = doc.add_paragraph()
code.add_run('''    @app.after_request
    def add_security_headers(response):
        response.headers['X-Content-Type-Options'] = 'nosniff'
        response.headers['X-Frame-Options'] = 'DENY'
        response.headers['Referrer-Policy'] = 'strict-origin-when-cross-origin'
        response.headers['Content-Security-Policy'] = "default-src 'self' data: blob:; frame-ancestors 'none';"
        return response''').font.name = 'Courier New'

p = doc.add_paragraph()
p.add_run('Rate Limiting & Protectie DDoS:').bold = True
p.add_run(' API-ul folosește Flask-Limiter pentru a restrictiona numarul de request-uri, blocand actiunile malitioase si tentativele de brute-force.')

# Section 5
doc.add_heading('5. Caching & Optimizare', level=1)
p = doc.add_paragraph()
p.add_run('Frontend Caching (localforage):').bold = True
p.add_run(' Se foloseste biblioteca localforage (o abstractizare peste IndexedDB) pentru a stoca local pachete de date. Aceasta abordare scade semnificativ numarul de request-uri catre server pentru datele care nu se schimba frecvent.')
p = doc.add_paragraph()
p.add_run('Optimizare Fisiere si Imagini:').bold = True
p.add_run(' Sistemul backend gestioneaza upload-urile prin generarea de fisiere sigure si suprascrierea structurata (ex. salvarea numelui original). Fisierele sunt servite eficient utilizand send_from_directory, care negociaza headere de caching cu browser-ul.')
p = doc.add_paragraph()
p.add_run('Optimizare Bundle (Vite):').bold = True
p.add_run(' La compilarea aplicatiei (vite build), codul JavaScript si CSS este minify-at si impartit in "chunks" (chunk splitting). Acest lucru reduce dramatic timpul de incarcare initiala (Time To Interactive / First Contentful Paint).')
p = doc.add_paragraph()
p.add_run('Optimizari Baze de Date:').bold = True
p.add_run(' Preluarea de liste lungi se face folosind paginare sau clauze LIMIT (de ex. LIMIT 15 pentru cartile recente), pastrand footprint-ul pe memorie si trafic cat mai mic.')

# Section 6
doc.add_heading('6. Scalabilitate', level=1)
doc.add_paragraph('Arhitectura este gandita pentru a scala orizontal extrem de rapid.')
p = doc.add_paragraph()
p.add_run('Stateless Backend:').bold = True
p.add_run(' Datorita folosirii JWT (JSON Web Tokens), serverul Flask nu pastreaza starea sesiunilor in memoria interna. Prin urmare, sistemul poate fi scalat lansand mai multe containere Flask / instante Gunicorn in spatele unui Load Balancer fara a necesita mecanisme complexe de "sticky sessions".')
p = doc.add_paragraph()
p.add_run('Nginx Reverse Proxy & Load Balancing:').bold = True
p.add_run(' Nginx actioneaza atat ca un web server ultrarapid pentru fisierele statice Vue (din folderul /dist), cat si ca reverse proxy, dirijand eficient traficul de API spre backend.')
code = doc.add_paragraph()
code.add_run('''    # Proxy API requests to Flask/Gunicorn
    location /api/ {
        proxy_pass         http://127.0.0.1:8000;
        proxy_set_header   Host $host;
        proxy_set_header   X-Real-IP $remote_addr;
        proxy_set_header   X-Forwarded-For $proxy_add_x_forwarded_for;
    }''').font.name = 'Courier New'

doc.save('C:/Users/zimbr/biblioteca-repo/Documentatie_Tehnica_Bibliosmart.docx')
print('Document generat cu succes!')
